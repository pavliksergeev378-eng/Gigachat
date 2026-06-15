# -*- coding: utf-8 -*-
"""
Генератор полной схемы алгоритма «Организация обращения» v2 в формате .docx.
A4 landscape. Требует: python-docx.

Запуск:
    pip install python-docx
    python generate-algorithm-doc.py
"""
import os
from docx import Document
from docx.shared import Mm, Pt, RGBColor, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Алгоритм-Организация-обращения.docx')

# ============ Цвета ============
C_HEADER = '2C3E50'
C_ACCENT = 'B8895D'
C_GREY_SOFT = 'F5F5F5'
C_GREY = 'E8E8E8'
C_GREY_DARK = '888888'
C_BLUE_SOFT = 'EAF2FA'
C_BLUE = 'D6E6F4'
C_BLUE_DARK = 'A7C7E7'
C_GREEN_SOFT = 'E8F5E9'
C_GREEN = 'C8E6C9'
C_GREEN_DARK = '81C784'
C_YELLOW_SOFT = 'FFF8E1'
C_YELLOW = 'FFE082'
C_YELLOW_DARK = 'FFB300'
C_RED_SOFT = 'FFEBEE'
C_RED = 'FFCDD2'
C_RED_DARK = 'E57373'
C_CODE_BG = 'F8F8F8'


# ============ Низкоуровневые помощники ============

def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_cell_border(cell, color='BBBBBB', size='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # удалить старые tcBorders если есть
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=80, bottom=80, left=140, right=140):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def add_run(par, text, *, bold=False, italic=False, size=11, color=None, font='Arial'):
    run = par.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_par(container, text='', *, size=11, bold=False, italic=False, color=None,
            align=None, space_before=0, space_after=4, font='Arial'):
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size, color=color, font=font)
    return p


def add_runs_par(container, runs, *, align=None, space_before=0, space_after=4):
    """runs: list of dicts {text, bold?, italic?, size?, color?}"""
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    for r in runs:
        add_run(
            p, r['text'],
            bold=r.get('bold', False),
            italic=r.get('italic', False),
            size=r.get('size', 11),
            color=r.get('color'),
            font=r.get('font', 'Arial'),
        )
    return p


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def fill_cell(cell, text, *, fill=None, bold=False, italic=False, size=10, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, border_color='BBBBBB', font='Arial', mono=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # очистить дефолтный пустой параграф
    if cell.paragraphs and not cell.paragraphs[0].text:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if isinstance(text, list):
        for i, line in enumerate(text):
            if i > 0:
                p = cell.add_paragraph()
                p.alignment = align
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
            if isinstance(line, dict):
                add_run(p, line['text'], bold=line.get('bold', bold),
                        italic=line.get('italic', italic),
                        size=line.get('size', size),
                        color=line.get('color', color),
                        font=line.get('font', font))
            else:
                add_run(p, line, bold=bold, italic=italic, size=size, color=color,
                        font='Consolas' if mono else font)
    else:
        add_run(p, text, bold=bold, italic=italic, size=size, color=color,
                font='Consolas' if mono else font)
    if fill:
        shade_cell(cell, fill)
    set_cell_border(cell, color=border_color)
    set_cell_margins(cell)


# ============ Сборка документа ============

def main():
    doc = Document()

    # A4 landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    # Базовый шрифт документа
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # =================== TITLE PAGE ===================
    for _ in range(3):
        add_par(doc, '', space_after=0)

    add_par(doc, 'АЛГОРИТМ', size=22, bold=True, color=C_HEADER,
            align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_par(doc, '«Организация обращения»', size=36, bold=True, color=C_ACCENT,
            align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_par(doc, 'Версия 2 · 4 фазы · 10 финальных исходов', size=14,
            color='666666', align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_par(doc, '', space_after=0)
    add_par(doc, 'Полная схема выполнения с пояснениями', size=14, italic=True,
            color='888888', align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    for _ in range(6):
        add_par(doc, '', space_after=0)
    add_par(doc, 'GigaChat-Platform · Дашборд алгоритмов', size=10,
            color='AAAAAA', align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    page_break(doc)

    # =================== 1. ОБЗОР ===================
    add_par(doc, 'Обзор', size=20, bold=True, color=C_HEADER, space_after=8)

    add_par(doc, 'Что делает алгоритм', size=14, bold=True, color='444444',
            space_before=4, space_after=4)
    add_par(doc,
            'Принимает служебную записку (Word или PDF), извлекает 4 идентификатора заявителя, '
            'затем последовательно проходит 4 фазы проверки в БД и возвращает один из 10 финальных исходов '
            'с пояснением, какие действия требуются для дальнейшей обработки обращения.',
            size=11, space_after=8)

    add_par(doc, 'Технический стек', size=14, bold=True, color='444444',
            space_before=8, space_after=4)

    stack_table = doc.add_table(rows=1, cols=4)
    stack_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    stack_table.autofit = False
    widths = [Mm(60), Mm(60), Mm(70), Mm(70)]
    rows = [
        ('Frontend', 'HTML / JS на дашборде', 'Drag-n-drop, JSZip для .docx', C_BLUE_SOFT),
        ('Webhook + оркестрация', 'n8n', '30 узлов, 22 связи, 4 фазы', C_BLUE_SOFT),
        ('Извлечение текста .pdf', 'OCR-сервис', 'localhost:8055/extract', C_BLUE_SOFT),
        ('Извлечение полей', 'GigaChat LLM (локальный)', 'OpenAI-совместимый API', C_BLUE_SOFT),
        ('БД', 'PostgreSQL', '3 таблицы + UNIQUE INDEX на ФИО', C_BLUE_SOFT),
    ]
    stack_table.rows[0].cells[0].text = ''
    # заголовок
    hdr = stack_table.rows[0]
    fill_cell(hdr.cells[0], 'Слой',        bold=True, fill=C_HEADER, color='FFFFFF', size=11)
    fill_cell(hdr.cells[1], 'Компонент',   bold=True, fill=C_HEADER, color='FFFFFF', size=11)
    fill_cell(hdr.cells[2], 'Назначение',  bold=True, fill=C_HEADER, color='FFFFFF', size=11)
    fill_cell(hdr.cells[3], 'Примечание',  bold=True, fill=C_HEADER, color='FFFFFF', size=11)
    for i, w in enumerate(widths):
        hdr.cells[i].width = w

    for layer, comp, purp, note_fill in rows:
        row = stack_table.add_row()
        fill_cell(row.cells[0], layer, bold=True, size=10)
        fill_cell(row.cells[1], comp, size=10)
        fill_cell(row.cells[2], purp, size=10, color='555555')
        fill_cell(row.cells[3], '', fill=note_fill, size=10)
        # перезаписать ячейку с note текстом (не fill)
        for i, w in enumerate(widths):
            row.cells[i].width = w
    # фикс последней колонки — текст про колонку Примечание
    notes_text = [
        'Поддержка .docx и .pdf, до 50 МБ',
        'Активируется в n8n через API-токен',
        'Возвращает JSON с полем text',
        'Модель локальная, токен не требуется',
        'NULLABLE employee_number — ключевая особенность',
    ]
    for i, txt in enumerate(notes_text, start=1):
        cell = stack_table.rows[i].cells[3]
        # очистить
        cell.text = ''
        fill_cell(cell, txt, size=10, color='555555')

    page_break(doc)

    # =================== 2. ВХОДНЫЕ ДАННЫЕ ===================
    add_par(doc, 'Входные данные: 4 идентификатора', size=20, bold=True, color=C_HEADER,
            space_after=8)
    add_par(doc,
            'Алгоритм извлекает из служебной записки следующие 4 поля. '
            'Личный номер ИЛИ СНИЛС достаточно одного — но без хотя бы одного из них документ не пройдёт Фазу 1.',
            size=11, space_after=10)

    fields_table = doc.add_table(rows=1, cols=4)
    fields_table.autofit = False
    widths = [Mm(50), Mm(50), Mm(80), Mm(80)]
    hdr = fields_table.rows[0]
    fill_cell(hdr.cells[0], 'Поле',         bold=True, fill=C_HEADER, color='FFFFFF', size=11)
    fill_cell(hdr.cells[1], 'Формат',       bold=True, fill=C_HEADER, color='FFFFFF', size=11)
    fill_cell(hdr.cells[2], 'Где в документе',bold=True, fill=C_HEADER, color='FFFFFF', size=11)
    fill_cell(hdr.cells[3], 'Обязательность',bold=True, fill=C_HEADER, color='FFFFFF', size=11)
    for i, w in enumerate(widths):
        hdr.cells[i].width = w

    fields_data = [
        ('ФИО', 'Фамилия Имя Отчество',
         'Строка после «От:»',
         'Обязательно. Без него — отказ на Фазе 1 (исход A)'),
        ('Дата рождения', 'ДД.ММ.ГГГГ',
         'Строка после «Дата рождения:»',
         'Обязательно. Без неё — отказ на Фазе 1 (исход B)'),
        ('Личный номер', 'ХХХ-ХХХ (6 цифр через дефис)',
         'Строка после «Личный номер:»',
         'Достаточно одного из двух (или ЛН, или СНИЛС)'),
        ('СНИЛС', 'ХХХ-ХХХ-ХХХ ХХ (3-3-3 дефис, пробел, 2)',
         'Строка после «СНИЛС:»',
         'Достаточно одного из двух'),
    ]
    for name, fmt, where, oblig in fields_data:
        row = fields_table.add_row()
        fill_cell(row.cells[0], name, bold=True, size=11)
        fill_cell(row.cells[1], fmt, size=10, mono=True, color='2C3E50')
        fill_cell(row.cells[2], where, size=10, color='555555')
        fill_cell(row.cells[3], oblig, size=10, color='555555')
        for i, w in enumerate(widths):
            row.cells[i].width = w

    add_par(doc, '', space_after=12)
    add_par(doc, 'Пример валидной служебной записки', size=13, bold=True,
            color='444444', space_before=8, space_after=4)
    sample = [
        'СЛУЖЕБНАЯ ЗАПИСКА',
        '',
        'От: Иванова Иванна Ивановна',
        'Дата рождения: 15.03.1985',
        'Личный номер: 100-100',
        'СНИЛС: 100-100-100 10',
        '',
        'Прошу рассмотреть моё обращение по поводу предоставления отпуска.',
    ]
    sample_table = doc.add_table(rows=1, cols=1)
    sample_table.autofit = False
    sample_table.rows[0].cells[0].width = Mm(200)
    fill_cell(sample_table.rows[0].cells[0], sample, fill=C_CODE_BG, size=10,
              mono=True, color='2C3E50', border_color='DDDDDD')

    page_break(doc)

    # =================== 3. ОБЩАЯ СХЕМА ВЫПОЛНЕНИЯ ===================
    add_par(doc, 'Общая схема выполнения', size=20, bold=True, color=C_HEADER,
            space_after=4)
    add_par(doc, 'Поток данных от загрузки документа до одного из 10 финальных исходов (A — J).',
            size=10, italic=True, color='888888', space_after=10)

    # Полная ширина в landscape: 297mm - 18*2 = 261mm
    TOTAL_W = 261
    flow = doc.add_table(rows=1, cols=3)
    flow.autofit = False

    def flow_row(t, rowspec):
        """rowspec: list of (text, fill, width_mm, opts={})"""
        row = t.add_row()
        # merge cells if fewer items than columns
        if len(rowspec) == 1:
            row.cells[0].merge(row.cells[-1])
            spec = rowspec[0]
            fill_cell(row.cells[0], spec[0], fill=spec[1], size=spec[3].get('size', 11),
                      bold=spec[3].get('bold', False), color=spec[3].get('color'),
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            row.cells[0].width = Mm(spec[2])
        elif len(rowspec) == 3:
            for i, spec in enumerate(rowspec):
                fill_cell(row.cells[i], spec[0], fill=spec[1], size=spec[3].get('size', 10),
                          bold=spec[3].get('bold', False), color=spec[3].get('color'),
                          align=WD_ALIGN_PARAGRAPH.CENTER)
                row.cells[i].width = Mm(spec[2])
        elif len(rowspec) == 2:
            row.cells[0].merge(row.cells[1])
            for i, spec in enumerate(rowspec):
                cell_idx = 0 if i == 0 else 2
                fill_cell(row.cells[cell_idx], spec[0], fill=spec[1],
                          size=spec[3].get('size', 10),
                          bold=spec[3].get('bold', False), color=spec[3].get('color'),
                          align=WD_ALIGN_PARAGRAPH.CENTER)
                row.cells[cell_idx].width = Mm(spec[2])

    # Заголовок таблицы — сразу первый шаг
    first = flow.rows[0]
    first.cells[0].merge(first.cells[2])
    fill_cell(first.cells[0], 'ВХОД: Word или PDF документ (служебная записка)',
              fill=C_BLUE_DARK, color='FFFFFF', bold=True, size=12,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    first.cells[0].width = Mm(TOTAL_W)

    # Стрелка вниз
    arrow_row = flow.add_row()
    arrow_row.cells[0].merge(arrow_row.cells[2])
    fill_cell(arrow_row.cells[0], '↓', size=14, bold=True, color=C_GREY_DARK,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    flow_row(flow, [('Извлечение текста: .docx → JSZip (браузер) | .pdf → OCR сервис',
                    C_BLUE, TOTAL_W, {'bold': True, 'size': 11})])

    arrow_row = flow.add_row()
    arrow_row.cells[0].merge(arrow_row.cells[2])
    fill_cell(arrow_row.cells[0], '↓', size=14, bold=True, color=C_GREY_DARK,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    flow_row(flow, [('LLM-парсинг: GigaChat извлекает 4 поля (ФИО, ДР, ЛН, СНИЛС)',
                    C_BLUE, TOTAL_W, {'bold': True, 'size': 11})])

    arrow_row = flow.add_row()
    arrow_row.cells[0].merge(arrow_row.cells[2])
    fill_cell(arrow_row.cells[0], '↓', size=14, bold=True, color=C_GREY_DARK,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    # Banner Phase 1
    flow_row(flow, [('ФАЗА 1 · Валидация полей документа',
                    C_HEADER, TOTAL_W, {'bold': True, 'size': 13, 'color': 'FFFFFF'})])

    third = TOTAL_W // 3
    rem = TOTAL_W - 2 * third
    flow_row(flow, [
        ('A · нет ФИО\n→ «ФИО заявителя отсутствует»', C_RED, third, {'size': 9, 'bold': True}),
        ('B · нет даты рождения\n→ «Дата рождения отсутствует»', C_RED, third, {'size': 9, 'bold': True}),
        ('C · нет и ЛН и СНИЛС\n→ «Личный номер или СНИЛС отсутствует»', C_RED, rem, {'size': 9, 'bold': True}),
    ])
    arrow_row = flow.add_row()
    arrow_row.cells[0].merge(arrow_row.cells[2])
    fill_cell(arrow_row.cells[0], '↓ все 4 поля валидны',
              size=10, italic=True, color=C_GREY_DARK,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    # Banner Phase 2
    flow_row(flow, [('ФАЗА 2 · Идентификация в реестре  ·  SELECT … FROM appeal_employees',
                    C_HEADER, TOTAL_W, {'bold': True, 'size': 13, 'color': 'FFFFFF'})])
    flow_row(flow, [
        ('D · ФИО не в реестре\n→ «Заявитель не идентифицирован»', C_RED, third, {'size': 9, 'bold': True}),
        ('E · в реестре, но employee_number = NULL\n→ «Сведений недостаточно»', C_YELLOW, third, {'size': 9, 'bold': True}),
        ('OK · найден + ТН\n→ переход к Фазе 3', C_GREEN, rem, {'size': 9, 'bold': True}),
    ])
    arrow_row = flow.add_row()
    arrow_row.cells[0].merge(arrow_row.cells[2])
    fill_cell(arrow_row.cells[0], '↓ ТН вписан в карточку',
              size=10, italic=True, color=C_GREY_DARK,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    # Banner Phase 3
    flow_row(flow, [('ФАЗА 3 · Мероприятие №1  ·  SELECT … FROM appeal_event1',
                    C_HEADER, TOTAL_W, {'bold': True, 'size': 13, 'color': 'FFFFFF'})])
    flow_row(flow, [
        ('F · нет записи в event1\n→ «Сведения отсутствуют»', C_RED, third, {'size': 9, 'bold': True}),
        ('G · is_done = FALSE\n→ «Мероприятие №1 не выполнено»', C_YELLOW, third, {'size': 9, 'bold': True}),
        ('OK · is_done = TRUE\n→ переход к Фазе 4', C_GREEN, rem, {'size': 9, 'bold': True}),
    ])
    arrow_row = flow.add_row()
    arrow_row.cells[0].merge(arrow_row.cells[2])
    fill_cell(arrow_row.cells[0], '↓',
              size=14, bold=True, color=C_GREY_DARK,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    # Banner Phase 4
    flow_row(flow, [('ФАЗА 4 · Мероприятие №2  ·  SELECT … FROM appeal_event2',
                    C_HEADER, TOTAL_W, {'bold': True, 'size': 13, 'color': 'FFFFFF'})])
    flow_row(flow, [
        ('H · нет записи в event2\n→ «Сведения отсутствуют»', C_RED, third, {'size': 9, 'bold': True}),
        ('I · is_done = FALSE\n→ «Необходимо обратиться за помощью»', C_YELLOW, third, {'size': 9, 'bold': True}),
        ('J · is_done = TRUE\n🎉 «Все мероприятия пройдены успешно»', C_GREEN_DARK, rem, {'size': 9, 'bold': True, 'color': 'FFFFFF'}),
    ])

    page_break(doc)

    # =================== 4. ФАЗА 1 ===================
    phase_section(doc, 1,
                  title='Фаза 1 · Валидация полей документа',
                  goal='Проверить, что LLM смогла извлечь все 4 обязательных поля (с учётом «или ЛН, или СНИЛС»).',
                  where='Узел «Фаза 1: проверка полей» в workflow (Code-нода).',
                  sql=None,
                  outcomes=[
                      ('A', C_RED_DARK, '✗', 'ФИО отсутствует',
                       'Поле fio пустое после парсинга LLM',
                       '«ФИО заявителя отсутствует, нужно предоставить данные»',
                       '01-no-fio.docx / .pdf'),
                      ('B', C_RED_DARK, '✗', 'Дата рождения отсутствует',
                       'Поле birth_date пустое',
                       '«Дата рождения заявителя отсутствует, нужно предоставить данные»',
                       '02-no-birth.docx / .pdf'),
                      ('C', C_RED_DARK, '✗', 'И ЛН, и СНИЛС отсутствуют',
                       'И personal_number, и snils пустые одновременно',
                       '«Личный номер или СНИЛС заявителя отсутствует, нужно предоставить данные»',
                       '03-no-id.docx / .pdf'),
                      ('OK', C_GREEN_DARK, '✓', 'Все поля валидны',
                       'fio + birth_date + (personal_number или snils) присутствуют',
                       '«Служебная записка идентифицирована» → переход к Фазе 2',
                       '— (продолжение)'),
                  ],
                  notes=[
                      'Проверка идёт по порядку: сначала ФИО, потом ДР, потом (ЛН ИЛИ СНИЛС). '
                      'Если ФИО пустое — сразу останов, остальные не проверяются.',
                      'Регекс-парсер допускает форматы «Личный номер: 123-456», «123456» (нормализуется в 123-456), '
                      'но отвергает форматы вроде «12-34-56» или текстовые значения.',
                  ])

    page_break(doc)

    # =================== 5. ФАЗА 2 ===================
    phase_section(doc, 2,
                  title='Фаза 2 · Идентификация в реестре',
                  goal='Найти ФИО заявителя в БД сотрудников и получить его табельный номер.',
                  where='Узел «БД №1: appeal_employees» (Postgres) + «Фаза 2: классификация» (Code).',
                  sql="SELECT full_name, employee_number FROM appeal_employees\n"
                      "WHERE full_name ILIKE $1 ESCAPE '\\'",
                  outcomes=[
                      ('D', C_RED_DARK, '✗', 'ФИО не в реестре',
                       'SQL-запрос вернул 0 строк',
                       '«Заявитель не идентифицирован, нужно предоставить дополнительные данные»',
                       '04-not-in-db (Кузнецов А. С.)'),
                      ('E', C_YELLOW_DARK, '!', 'В реестре, но без табельного',
                       'Найдена строка, но employee_number = NULL',
                       '«Сведений о заявителе недостаточно для идентификации, нужно сделать запрос»',
                       '05-no-tab-number (Сидоров С. С.)'),
                      ('OK', C_GREEN_DARK, '✓', 'ФИО + табельный номер найдены',
                       'employee_number — непустое значение',
                       'ТН вписывается в карточку, переход к Фазе 3',
                       '— (продолжение)'),
                  ],
                  notes=[
                      'NULLABLE-колонка employee_number — ключевая особенность v2: '
                      'позволяет различать «не знаем такого сотрудника» (D) и «знаем, но ТН ещё не присвоен» (E).',
                      'ILIKE с ESCAPE — case-insensitive точное совпадение (без wildcards). '
                      'fio_for_sql экранирует %, _, \\ для защиты от LIKE-инъекций.',
                      'При отказе БД (после фикса 1 из аудита) выдаётся явное сообщение об ошибке, а не псевдо-D.',
                  ])

    page_break(doc)

    # =================== 6. ФАЗА 3 ===================
    phase_section(doc, 3,
                  title='Фаза 3 · Проверка Мероприятия №1',
                  goal='Узнать, выполнено ли первое мероприятие для данного сотрудника.',
                  where='Узел «БД №2: appeal_event1» (Postgres) + «Фаза 3: классификация» (Code).',
                  sql="SELECT full_name, is_done FROM appeal_event1\n"
                      "WHERE full_name ILIKE $1 ESCAPE '\\'",
                  outcomes=[
                      ('F', C_RED_DARK, '✗', 'Нет записи в event1',
                       'SQL-запрос вернул 0 строк',
                       '«Сведения по данному заявителю отсутствуют»',
                       '06-not-in-event1 (Богданов Б. Б.)'),
                      ('G', C_YELLOW_DARK, '!', 'is_done = FALSE',
                       'Запись есть, но мероприятие не выполнено',
                       '«Для данного заявителя мероприятие №1 не выполнено»',
                       '07-event1-not-done (Морозов М. М.)'),
                      ('OK', C_GREEN_DARK, '✓', 'Мероприятие №1 выполнено',
                       'is_done = TRUE',
                       '«Мероприятие №1 для заявителя выполнено» → переход к Фазе 4',
                       '— (продолжение)'),
                  ],
                  notes=[
                      'Несмотря на UNIQUE INDEX на full_name, код использует rows.some(r => r.is_done === true) '
                      'для устойчивости к дубликатам (если бы они появились).',
                  ])

    page_break(doc)

    # =================== 7. ФАЗА 4 ===================
    phase_section(doc, 4,
                  title='Фаза 4 · Проверка Мероприятия №2 (финальная)',
                  goal='Узнать, выполнено ли второе мероприятие. Это последняя фаза — её исход и есть итог алгоритма.',
                  where='Узел «БД №3: appeal_event2» (Postgres) + «Фаза 4: классификация» (Code).',
                  sql="SELECT full_name, is_done FROM appeal_event2\n"
                      "WHERE full_name ILIKE $1 ESCAPE '\\'",
                  outcomes=[
                      ('H', C_RED_DARK, '✗', 'Нет записи в event2',
                       'SQL-запрос вернул 0 строк',
                       '«Сведения по данному заявителю отсутствуют»',
                       '08-not-in-event2 (Михайлов Д. А.)'),
                      ('I', C_YELLOW_DARK, '!', 'is_done = FALSE',
                       'Запись есть, но мероприятие не выполнено',
                       '«Данному заявителю необходимо обратиться за помощью»',
                       '09-event2-not-done (Васильев А. С.)'),
                      ('J', C_GREEN_DARK, '✓', 'Все мероприятия пройдены',
                       'is_done = TRUE',
                       '🎉 «Для данного заявителя все мероприятия пройдены успешно»',
                       '10-full-success (Иванова И. И.)'),
                  ],
                  notes=[
                      'У этой фазы нет IF-гейта после неё: все три исхода идут в один respondToWebhook. '
                      'Окрашивание итоговой плашки в HTML выбирается по статусу step4 (ok/not_done/not_found).',
                      'Исход J — единственный «зелёный» во всём алгоритме. Все остальные требуют какого-то действия.',
                  ])

    page_break(doc)

    # =================== 8. СВОДНАЯ ТАБЛИЦА 10 ИСХОДОВ ===================
    add_par(doc, 'Сводная таблица 10 финальных исходов', size=20, bold=True, color=C_HEADER,
            space_after=4)
    add_par(doc, 'Один файл на каждую ветку алгоритма. Полный прогон — 20 запусков (10 docx + 10 pdf).',
            size=10, italic=True, color='888888', space_after=10)

    sum_table = doc.add_table(rows=1, cols=6)
    sum_table.autofit = False
    widths_mm = [10, 45, 25, 45, 90, 46]
    hdr = sum_table.rows[0]
    headers = ['№', 'Файл', 'Фаза', 'Условие', 'Итоговое сообщение пользователю', 'Тестовое ФИО']
    for i, txt in enumerate(headers):
        fill_cell(hdr.cells[i], txt, bold=True, fill=C_HEADER, color='FFFFFF', size=11,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        hdr.cells[i].width = Mm(widths_mm[i])

    outcomes_full = [
        ('A', '01-no-fio',          'Фаза 1', 'fio пустое',
         'ФИО заявителя отсутствует, нужно предоставить данные',         '— (нет ФИО)',                C_RED_SOFT),
        ('B', '02-no-birth',        'Фаза 1', 'birth_date пустое',
         'Дата рождения заявителя отсутствует, нужно предоставить данные','Иванов Иван Иванович',       C_RED_SOFT),
        ('C', '03-no-id',           'Фаза 1', 'и ЛН, и СНИЛС пустые',
         'Личный номер или СНИЛС заявителя отсутствует',                  'Петров Пётр Петрович',       C_RED_SOFT),
        ('D', '04-not-in-db',       'Фаза 2', 'ФИО не в appeal_employees',
         'Заявитель не идентифицирован, нужно предоставить дополнительные данные','Кузнецов Алексей Сергеевич', C_RED_SOFT),
        ('E', '05-no-tab-number',   'Фаза 2', 'employee_number = NULL',
         'Сведений о заявителе недостаточно для идентификации, нужно сделать запрос','Сидоров Сидор Сидорович', C_YELLOW_SOFT),
        ('F', '06-not-in-event1',   'Фаза 3', 'нет записи в appeal_event1',
         'Сведения по данному заявителю отсутствуют',                     'Богданов Богдан Богданович', C_RED_SOFT),
        ('G', '07-event1-not-done', 'Фаза 3', 'event1.is_done = FALSE',
         'Для данного заявителя мероприятие №1 не выполнено',             'Морозов Михаил Михайлович',   C_YELLOW_SOFT),
        ('H', '08-not-in-event2',   'Фаза 4', 'нет записи в appeal_event2',
         'Сведения по данному заявителю отсутствуют',                     'Михайлов Дмитрий Александрович', C_RED_SOFT),
        ('I', '09-event2-not-done', 'Фаза 4', 'event2.is_done = FALSE',
         'Данному заявителю необходимо обратиться за помощью',            'Васильев Андрей Сергеевич',   C_YELLOW_SOFT),
        ('J', '10-full-success',    'Фаза 4', 'event2.is_done = TRUE',
         '🎉 Для данного заявителя все мероприятия пройдены успешно',     'Иванова Иванна Ивановна',     C_GREEN_SOFT),
    ]
    for letter, fname, phase, cond, msg, fio, fill in outcomes_full:
        row = sum_table.add_row()
        fill_cell(row.cells[0], letter, bold=True, fill=fill, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(row.cells[1], fname, fill=fill, size=9, mono=True, color='2C3E50')
        fill_cell(row.cells[2], phase, fill=fill, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(row.cells[3], cond, fill=fill, size=9, color='555555', mono=True)
        fill_cell(row.cells[4], msg, fill=fill, size=10)
        fill_cell(row.cells[5], fio, fill=fill, size=9, italic=True, color='555555')
        for i, w in enumerate(widths_mm):
            row.cells[i].width = Mm(w)

    page_break(doc)

    # =================== 9. ТЕСТИРОВАНИЕ ===================
    add_par(doc, 'Как тестировать алгоритм', size=20, bold=True, color=C_HEADER,
            space_after=8)

    add_par(doc, 'Перед запуском', size=14, bold=True, color='444444', space_before=4, space_after=4)
    steps = [
        '1. Раскатить SQL из База данных/OrgAppeal-Setup.md в Postgres (3 таблицы + ~130 строк тест-данных)',
        '2. Импортировать Workflow/organization-appeal.json в n8n и привязать credentials (Postgres + GigaChat)',
        '3. Активировать workflow в n8n',
        '4. Убедиться, что OCR-сервис слушает http://130.100.94.119:8055/extract',
        '5. Убедиться, что GigaChat LLM запущен и доступен с n8n-машины',
    ]
    for step in steps:
        add_par(doc, step, size=11, space_after=3)

    add_par(doc, 'Прогон тестов', size=14, bold=True, color='444444',
            space_before=10, space_after=4)
    add_par(doc, 'Открыть GigaChat-Platform.html → «Организация обращения» и по очереди перетащить все 20 файлов:',
            size=11, space_after=6)

    tests_table = doc.add_table(rows=1, cols=2)
    tests_table.autofit = False
    hdr = tests_table.rows[0]
    fill_cell(hdr.cells[0], 'Формат', bold=True, fill=C_HEADER, color='FFFFFF', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(hdr.cells[1], 'Что проверяет', bold=True, fill=C_HEADER, color='FFFFFF', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    hdr.cells[0].width = Mm(40)
    hdr.cells[1].width = Mm(220)

    test_rows = [
        ('10 × .docx', 'Браузерный путь: JSZip извлекает текст в браузере и отправляет JSON-ом в n8n. '
                       'OCR-сервис не задействован.'),
        ('10 × .pdf',  'Серверный путь: файл уходит в OCR localhost:8055/extract, затем — тот же LLM-парсинг.'),
    ]
    for fmt, what in test_rows:
        row = tests_table.add_row()
        fill_cell(row.cells[0], fmt, bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, fill=C_BLUE_SOFT)
        fill_cell(row.cells[1], what, size=10, color='444444')
        row.cells[0].width = Mm(40)
        row.cells[1].width = Mm(220)

    add_par(doc, '', space_after=8)
    add_par(doc, 'Каждый из 10 сценариев должен дать одинаковый итог в обоих форматах. '
            'Если итог расходится — баг в одной из веток (JSZip или OCR). См. Tests/OrgAppeal/README.md.',
            size=11, italic=True, color='666666', space_after=4)

    # =================== СОХРАНИТЬ ===================
    doc.save(OUTPUT)
    print(f'[ok] {OUTPUT}')
    print(f'     {os.path.getsize(OUTPUT) / 1024:.1f} КБ')


def phase_section(doc, n, *, title, goal, where, sql, outcomes, notes):
    """Шаблон страницы для каждой фазы."""
    add_par(doc, title, size=20, bold=True, color=C_HEADER, space_after=8)

    # Goal + Where
    info_table = doc.add_table(rows=2, cols=2)
    info_table.autofit = False
    fill_cell(info_table.rows[0].cells[0], 'Цель', bold=True, fill=C_GREY, size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(info_table.rows[0].cells[1], goal, size=11)
    fill_cell(info_table.rows[1].cells[0], 'Где в workflow', bold=True, fill=C_GREY, size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(info_table.rows[1].cells[1], where, size=10, color='555555')
    info_table.rows[0].cells[0].width = Mm(50)
    info_table.rows[0].cells[1].width = Mm(210)
    info_table.rows[1].cells[0].width = Mm(50)
    info_table.rows[1].cells[1].width = Mm(210)

    add_par(doc, '', space_after=6)

    # SQL
    if sql:
        add_par(doc, 'SQL-запрос', size=13, bold=True, color='444444',
                space_before=6, space_after=4)
        sql_table = doc.add_table(rows=1, cols=1)
        sql_table.autofit = False
        sql_table.rows[0].cells[0].width = Mm(260)
        fill_cell(sql_table.rows[0].cells[0], sql, fill=C_CODE_BG, size=10,
                  mono=True, color='2C3E50', border_color='DDDDDD')
        add_par(doc, '', space_after=4)

    # Outcomes
    add_par(doc, 'Возможные исходы', size=13, bold=True, color='444444',
            space_before=8, space_after=4)
    out_table = doc.add_table(rows=1, cols=6)
    out_table.autofit = False
    widths_mm = [12, 12, 50, 60, 90, 38]
    hdr = out_table.rows[0]
    headers = ['Метка', '', 'Условие', 'Детали', 'Сообщение пользователю', 'Тест-файл']
    for i, txt in enumerate(headers):
        fill_cell(hdr.cells[i], txt, bold=True, fill=C_HEADER, color='FFFFFF', size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        hdr.cells[i].width = Mm(widths_mm[i])

    for letter, color, glyph, cond, detail, msg, tf in outcomes:
        row = out_table.add_row()
        fill_cell(row.cells[0], letter, bold=True, fill=color, color='FFFFFF', size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(row.cells[1], glyph, bold=True, fill=color, color='FFFFFF', size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(row.cells[2], cond, bold=True, size=11)
        fill_cell(row.cells[3], detail, size=10, color='555555')
        fill_cell(row.cells[4], msg, size=10)
        fill_cell(row.cells[5], tf, size=9, italic=True, color='666666', mono=True)
        for i, w in enumerate(widths_mm):
            row.cells[i].width = Mm(w)

    # Notes
    if notes:
        add_par(doc, '', space_after=4)
        add_par(doc, 'Примечания', size=13, bold=True, color='444444',
                space_before=6, space_after=4)
        for note in notes:
            add_runs_par(doc, [
                {'text': '·  ', 'bold': True, 'size': 11, 'color': C_ACCENT},
                {'text': note, 'size': 10, 'color': '555555'},
            ], space_after=3)


if __name__ == '__main__':
    main()

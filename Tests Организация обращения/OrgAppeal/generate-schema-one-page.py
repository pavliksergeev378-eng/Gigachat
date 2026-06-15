# -*- coding: utf-8 -*-
"""
Одностраничная функциональная схема алгоритма «Организация обращения» v2.
A4 landscape, строго на одну страницу.

Запуск:
    pip install python-docx
    python generate-schema-one-page.py
"""
import os
from docx import Document
from docx.shared import Mm, Pt, RGBColor, Emu
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                     'Схема-Организация-обращения.docx')

# ============ Цвета ============
C_HEADER = '2C3E50'
C_ACCENT = 'B8895D'
C_BLUE_DARK = '6FA8DC'
C_BLUE = 'CFE2F3'
C_BLUE_SOFT = 'EAF2FA'
C_GREEN = 'C8E6C9'
C_GREEN_DARK = '548F5C'
C_YELLOW = 'FFE082'
C_YELLOW_DARK = 'B7791F'
C_RED = 'FFCDD2'
C_RED_DARK = 'C0392B'
C_GREEN_TEXT = '1B5E20'
C_BROWN_TEXT = '5D4037'


# ============ XML-помощники ============

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def cell_border(cell, color='AAAAAA', size='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tcb = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tcb.append(b)
    tcPr.append(tcb)


def cell_margins(cell, t=20, b=20, l=80, r=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    m = OxmlElement('w:tcMar')
    for edge, val in (('top', t), ('bottom', b), ('left', l), ('right', r)):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')
        m.append(e)
    tcPr.append(m)


def fill_cell(cell, lines, *, bg=None, fg='000000', bold=False, size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER, mono=False, border_color='AAAAAA',
              cell_t=20, cell_b=20):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if not isinstance(lines, list):
        lines = [lines]
    # first paragraph
    p0 = cell.paragraphs[0]
    p0.alignment = align
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(lines):
        if i == 0:
            p = p0
        else:
            p = cell.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        if isinstance(line, dict):
            r = p.add_run(line['text'])
            r.font.name = 'Consolas' if line.get('mono', mono) else 'Arial'
            r.font.size = Pt(line.get('size', size))
            r.bold = line.get('bold', bold)
            r.italic = line.get('italic', False)
            r.font.color.rgb = RGBColor.from_string(line.get('color', fg))
        else:
            r = p.add_run(line)
            r.font.name = 'Consolas' if mono else 'Arial'
            r.font.size = Pt(size)
            r.bold = bold
            r.font.color.rgb = RGBColor.from_string(fg)
    if bg:
        shade(cell, bg)
    cell_border(cell, color=border_color)
    cell_margins(cell, t=cell_t, b=cell_b)


def add_run(par, text, *, size=10, bold=False, italic=False, color='000000', font='Arial'):
    r = par.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)
    return r


def main():
    doc = Document()

    # Page setup: A4 landscape, тесные поля
    s = doc.sections[0]
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width = Mm(297)
    s.page_height = Mm(210)
    s.left_margin = Mm(10)
    s.right_margin = Mm(10)
    s.top_margin = Mm(8)
    s.bottom_margin = Mm(8)

    # Default style — мелкий шрифт, плотный интервал
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # ====== Заголовок ======
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_run(p, 'Функциональная схема алгоритма «Организация обращения»',
            size=15, bold=True, color=C_HEADER)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_run(p, 'v2  ·  4 фазы  ·  10 финальных исходов  ·  Word/PDF → 4 поля → проверка в 3 таблицах БД',
            size=9, italic=True, color='666666')

    # ====== Основная таблица: 3 колонки ======
    table = doc.add_table(rows=0, cols=3)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Полная ширина: 297 - 20 = 277 мм
    THIRD = Mm(92)
    FULL = Mm(277)

    def set_row_height(row, mm_val):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        # mm → DXA: 1mm = 56.7 DXA
        dxa = int(mm_val * 56.7)
        trHeight.set(qn('w:val'), str(dxa))
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)

    def full_row(content, *, bg, fg='FFFFFF', size=11, bold=True, height_mm=None):
        row = table.add_row()
        row.cells[0].merge(row.cells[2])
        fill_cell(row.cells[0], content, bg=bg, fg=fg, bold=bold, size=size)
        row.cells[0].width = FULL
        if height_mm:
            set_row_height(row, height_mm)
        return row

    def three_row(items, *, height_mm=None):
        """items: 3 dicts {lines, bg, fg}"""
        row = table.add_row()
        for i, it in enumerate(items):
            fill_cell(row.cells[i], it['lines'], bg=it['bg'], fg=it.get('fg', '000000'),
                      bold=False, size=10)
            row.cells[i].width = THIRD
        if height_mm:
            set_row_height(row, height_mm)
        return row

    # ====== ВХОД ======
    full_row('ВХОД: служебная записка (Word .docx или PDF)',
             bg=C_BLUE_DARK, size=12, height_mm=8)

    # ====== ИЗВЛЕЧЕНИЕ ======
    full_row([
        {'text': 'Извлечение текста', 'bold': True, 'size': 10, 'color': C_HEADER},
        {'text': '.docx → JSZip в браузере     ·     .pdf → OCR-сервис :8055/extract',
         'size': 9, 'color': '555555'},
    ], bg=C_BLUE, fg=C_HEADER, bold=False, size=10, height_mm=11)

    # ====== LLM ======
    full_row([
        {'text': 'LLM-парсинг (GigaChat)', 'bold': True, 'size': 10, 'color': C_HEADER},
        {'text': 'извлекает 4 поля: ФИО · Дата рождения · Личный номер (ХХХ-ХХХ) · СНИЛС (ХХХ-ХХХ-ХХХ ХХ)',
         'size': 9, 'color': '555555'},
    ], bg=C_BLUE, fg=C_HEADER, bold=False, size=10, height_mm=11)

    # ====== ФАЗА 1 ======
    full_row('ФАЗА 1  ·  Валидация полей документа',
             bg=C_HEADER, size=11, height_mm=6)
    three_row([
        {'lines': [
            {'text': 'A', 'bold': True, 'size': 18, 'color': 'FFFFFF'},
            {'text': 'нет ФИО', 'bold': True, 'size': 10, 'color': 'FFFFFF'},
            {'text': '«ФИО заявителя отсутствует»', 'size': 8, 'color': 'FFFFFF'},
        ], 'bg': C_RED_DARK, 'fg': 'FFFFFF'},
        {'lines': [
            {'text': 'B', 'bold': True, 'size': 18, 'color': 'FFFFFF'},
            {'text': 'нет даты рождения', 'bold': True, 'size': 10, 'color': 'FFFFFF'},
            {'text': '«Дата рождения отсутствует»', 'size': 8, 'color': 'FFFFFF'},
        ], 'bg': C_RED_DARK, 'fg': 'FFFFFF'},
        {'lines': [
            {'text': 'C', 'bold': True, 'size': 18, 'color': 'FFFFFF'},
            {'text': 'нет ЛН и СНИЛС', 'bold': True, 'size': 10, 'color': 'FFFFFF'},
            {'text': '«Личный номер или СНИЛС отсутствует»', 'size': 8, 'color': 'FFFFFF'},
        ], 'bg': C_RED_DARK, 'fg': 'FFFFFF'},
    ], height_mm=15)

    # ====== ФАЗА 2 ======
    full_row('ФАЗА 2  ·  Идентификация в реестре   |   SELECT … FROM appeal_employees',
             bg=C_HEADER, size=11, height_mm=6)
    three_row([
        {'lines': [
            {'text': 'D', 'bold': True, 'size': 18, 'color': 'FFFFFF'},
            {'text': 'ФИО не в реестре', 'bold': True, 'size': 10, 'color': 'FFFFFF'},
            {'text': '«Заявитель не идентифицирован»', 'size': 8, 'color': 'FFFFFF'},
        ], 'bg': C_RED_DARK, 'fg': 'FFFFFF'},
        {'lines': [
            {'text': 'E', 'bold': True, 'size': 18, 'color': C_BROWN_TEXT},
            {'text': 'employee_number = NULL', 'bold': True, 'size': 10, 'color': C_BROWN_TEXT},
            {'text': '«Сведений недостаточно»', 'size': 8, 'color': C_BROWN_TEXT},
        ], 'bg': C_YELLOW, 'fg': C_BROWN_TEXT},
        {'lines': [
            {'text': '→', 'bold': True, 'size': 18, 'color': C_GREEN_TEXT},
            {'text': 'ТН получен', 'bold': True, 'size': 10, 'color': C_GREEN_TEXT},
            {'text': 'продолжаем к Фазе 3', 'size': 8, 'color': C_GREEN_TEXT},
        ], 'bg': C_GREEN, 'fg': C_GREEN_TEXT},
    ], height_mm=15)

    # ====== ФАЗА 3 ======
    full_row('ФАЗА 3  ·  Мероприятие №1   |   SELECT … FROM appeal_event1',
             bg=C_HEADER, size=11, height_mm=6)
    three_row([
        {'lines': [
            {'text': 'F', 'bold': True, 'size': 18, 'color': 'FFFFFF'},
            {'text': 'нет записи в event1', 'bold': True, 'size': 10, 'color': 'FFFFFF'},
            {'text': '«Сведения отсутствуют»', 'size': 8, 'color': 'FFFFFF'},
        ], 'bg': C_RED_DARK, 'fg': 'FFFFFF'},
        {'lines': [
            {'text': 'G', 'bold': True, 'size': 18, 'color': C_BROWN_TEXT},
            {'text': 'is_done = FALSE', 'bold': True, 'size': 10, 'color': C_BROWN_TEXT},
            {'text': '«Мероприятие №1 не выполнено»', 'size': 8, 'color': C_BROWN_TEXT},
        ], 'bg': C_YELLOW, 'fg': C_BROWN_TEXT},
        {'lines': [
            {'text': '→', 'bold': True, 'size': 18, 'color': C_GREEN_TEXT},
            {'text': 'is_done = TRUE', 'bold': True, 'size': 10, 'color': C_GREEN_TEXT},
            {'text': 'продолжаем к Фазе 4', 'size': 8, 'color': C_GREEN_TEXT},
        ], 'bg': C_GREEN, 'fg': C_GREEN_TEXT},
    ], height_mm=15)

    # ====== ФАЗА 4 (финальная) ======
    full_row('ФАЗА 4  ·  Мероприятие №2 (финальная)   |   SELECT … FROM appeal_event2',
             bg=C_HEADER, size=11, height_mm=6)
    three_row([
        {'lines': [
            {'text': 'H', 'bold': True, 'size': 18, 'color': 'FFFFFF'},
            {'text': 'нет записи в event2', 'bold': True, 'size': 10, 'color': 'FFFFFF'},
            {'text': '«Сведения отсутствуют»', 'size': 8, 'color': 'FFFFFF'},
        ], 'bg': C_RED_DARK, 'fg': 'FFFFFF'},
        {'lines': [
            {'text': 'I', 'bold': True, 'size': 18, 'color': C_BROWN_TEXT},
            {'text': 'is_done = FALSE', 'bold': True, 'size': 10, 'color': C_BROWN_TEXT},
            {'text': '«Необходимо обратиться за помощью»', 'size': 8, 'color': C_BROWN_TEXT},
        ], 'bg': C_YELLOW, 'fg': C_BROWN_TEXT},
        {'lines': [
            {'text': 'J  ✓', 'bold': True, 'size': 18, 'color': 'FFFFFF'},
            {'text': 'is_done = TRUE', 'bold': True, 'size': 10, 'color': 'FFFFFF'},
            {'text': '«Все мероприятия пройдены успешно»', 'size': 8, 'color': 'FFFFFF'},
        ], 'bg': C_GREEN_DARK, 'fg': 'FFFFFF'},
    ], height_mm=15)

    # ====== Легенда / пояснения ======
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, '■ ', size=10, bold=True, color=C_RED_DARK)
    add_run(p, 'отказ (данных нет — нужны действия)   ', size=8, color='555555')
    add_run(p, '■ ', size=10, bold=True, color=C_YELLOW_DARK)
    add_run(p, 'данные есть, мероприятие не завершено   ', size=8, color='555555')
    add_run(p, '■ ', size=10, bold=True, color=C_GREEN_DARK)
    add_run(p, 'успех — переход дальше / финал   ', size=8, color='555555')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_run(p,
            'Алгоритм останавливается на первом несоблюдённом условии и возвращает соответствующее сообщение пользователю. '
            'Прохождение всех 4 фаз без отказа — единственный путь к финальному «успешно».',
            size=8, italic=True, color='888888')

    # Save
    doc.save(OUTPUT)
    print(f'[ok] {OUTPUT}')
    print(f'     {os.path.getsize(OUTPUT) / 1024:.1f} КБ')


if __name__ == '__main__':
    main()

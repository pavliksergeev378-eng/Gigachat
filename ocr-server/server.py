"""
Локальный HTTP-сервис извлечения текста из документов для GigaChat.

Делает то же, что коммерческий OCR-as-a-service, но полностью офлайн:
  - PDF: сначала пробует прямое извлечение через PyMuPDF (быстро, точно).
         Если в PDF мало текста — рендерит страницы как картинки и пропускает
         через EasyOCR (для сканированных документов).
  - DOCX: прямое извлечение через python-docx.
  - Картинки (PNG/JPG/TIFF/BMP/WEBP): EasyOCR.
  - TXT/MD/CSV: просто читаем как UTF-8.

n8n-workflow «organization-appeal» (и другие, использующие OCR) шлют файлы
на эндпоинт POST /extract с полем multipart `file` и ожидают JSON
с полем `text`.

Запуск:
    pip install -r requirements.txt
    python server.py

Конфигурация через переменные окружения:
    OCR_HOST           — на каком интерфейсе слушать. По умолчанию: 0.0.0.0
    OCR_PORT           — порт. По умолчанию: 8055
    OCR_LANGS          — языки для EasyOCR. По умолчанию: ru,en
    OCR_PDF_MIN_TEXT   — если PDF после PyMuPDF дал меньше N символов —
                          fallback на EasyOCR. По умолчанию: 50
    OCR_PDF_MAX_PAGES  — лимит страниц PDF для OCR (защита от мега-документов).
                          По умолчанию: 50
    OCR_EASYOCR_DIR    — папка с моделями EasyOCR (см. README). Если не задано —
                          ~/.EasyOCR (стандартное место).
    OCR_DEVICE         — cpu / cuda. По умолчанию: auto.

API:
    GET  /status
        resp: {"status": "ok", "easyocr_ready": bool, "device": "cpu|cuda"}

    POST /extract
        multipart: file=<binary>
        resp: {"text": "...", "source": "pymupdf|pymupdf+easyocr|docx|easyocr|plain", "filename": "..."}
        Универсальный эндпоинт: PDF, DOCX, картинки, TXT — определяется по расширению.
"""

import io
import os
import logging
import tempfile
from typing import Optional, TYPE_CHECKING

if TYPE_CHECKING:
    from typing import Any

# Принудительный оффлайн-режим: EasyOCR/torch не должны лезть в интернет
# даже если detect что моделей нет. Безопаснее, чем дефолт.
os.environ.setdefault('HF_HUB_OFFLINE', '1')
os.environ.setdefault('TRANSFORMERS_OFFLINE', '1')

import uvicorn
from fastapi import FastAPI, UploadFile, File, HTTPException

# OpenCV импортируется один раз глобально (используется в детекции таблиц)
import cv2

HOST = os.environ.get('OCR_HOST', '0.0.0.0')
PORT = int(os.environ.get('OCR_PORT', '8055'))
LANGS = os.environ.get('OCR_LANGS', 'ru,en').split(',')
PDF_MIN_TEXT = int(os.environ.get('OCR_PDF_MIN_TEXT', '50'))
PDF_MAX_PAGES = int(os.environ.get('OCR_PDF_MAX_PAGES', '50'))
MAX_UPLOAD_MB = int(os.environ.get('OCR_MAX_UPLOAD_MB', '50'))
EASYOCR_DIR = os.environ.get('OCR_EASYOCR_DIR', '').strip() or None
DEVICE = os.environ.get('OCR_DEVICE', '').strip().lower()

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
)
log = logging.getLogger('ocr-server')


# --- Ленивая инициализация EasyOCR ---
# EasyOCR грузит модели на ~150 МБ — делаем это при первом обращении,
# чтобы /health отвечал моментально (workflow его пингует часто).
_easyocr_reader = None
_easyocr_init_error: Optional[str] = None


def _detect_device() -> str:
    if DEVICE in ('cpu', 'cuda'):
        return DEVICE
    try:
        import torch
        if torch.cuda.is_available():
            return 'cuda'
    except Exception:
        pass
    return 'cpu'


def get_easyocr():
    """Инициализирует EasyOCR при первом вызове; повторно — возвращает кеш."""
    global _easyocr_reader, _easyocr_init_error
    if _easyocr_reader is not None:
        return _easyocr_reader
    # L6: НЕ блокируем повторные попытки навсегда — разовый сбой (например,
    # модели ещё не успели скопироваться) не должен требовать рестарта сервиса.
    # Пробуем инициализацию заново при каждом вызове, пока reader не готов.
    try:
        import easyocr
        use_gpu = _detect_device() == 'cuda'
        log.info('Loading EasyOCR (langs=%s, gpu=%s, dir=%s)', LANGS, use_gpu, EASYOCR_DIR or 'default')
        _easyocr_reader = easyocr.Reader(
            LANGS,
            gpu=use_gpu,
            model_storage_directory=EASYOCR_DIR,
            download_enabled=False,  # КРИТИЧНО: офлайн. Если моделей нет — упадёт явно.
            verbose=False,
        )
        _easyocr_init_error = None  # L6: сбрасываем прошлую ошибку после успеха
        log.info('EasyOCR ready.')
        return _easyocr_reader
    except Exception as e:
        _easyocr_init_error = (
            f'EasyOCR недоступен: {e}. '
            f'Проверь, что модели лежат в {EASYOCR_DIR or "~/.EasyOCR/model/"} '
            'и EASYOCR_DIR указан корректно.'
        )
        log.error(_easyocr_init_error)
        raise RuntimeError(_easyocr_init_error)


# --- Детекция таблиц через OpenCV ---

def _preprocess_for_lines(img: np.ndarray) -> np.ndarray:
    """
    Препроцессинг изображения для поиска линий:
    Grayscale → GaussianBlur → AdaptiveThreshold → BitwiseNot.
    """
    gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
    blurred = cv2.GaussianBlur(gray, (5, 5), 0)
    thresh = cv2.adaptiveThreshold(
        blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY, 11, 2
    )
    return cv2.bitwise_not(thresh)


def _find_lines(binary: np.ndarray) -> tuple[list[int], list[int]]:
    """
    Находит горизонтальные и вертикальные линии на бинарном изображении
    через Canny + HoughLinesP.
    Возвращает (y_coords, x_coords) — отсортированные уникальные координаты линий.
    """
    edges = cv2.Canny(binary, 50, 150, apertureSize=3)

    # Морфология: усиливаем линии
    kernel_h = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
    kernel_v = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
    edges_h = cv2.morphologyEx(edges, cv2.MORPH_CLOSE, kernel_h)
    edges_v = cv2.morphologyEx(edges, cv2.MORPH_CLOSE, kernel_v)
    edges = cv2.bitwise_or(edges_h, edges_v)

    lines = cv2.HoughLinesP(
        edges, rho=1, theta=np.pi / 180,
        threshold=100, minLineLength=80, maxLineGap=10
    )
    if lines is None:
        return [], []

    h_coords: set[int] = set()
    v_coords: set[int] = set()
    h, w = binary.shape[:2]
    margin = max(1, int(min(h, w) * 0.01))  # 1% от края

    for line in lines:
        x1, y1, x2, y2 = line[0]
        dx = abs(x2 - x1)
        dy = abs(y2 - y1)
        length = max(dx, dy)
        if length < 40:
            continue

        # Горизонтальная линия (угол ~0°)
        if dy < dx * 0.2:
            y_avg = (y1 + y2) // 2
            if margin <= y_avg <= h - margin:
                h_coords.add(y_avg)
        # Вертикальная линия (угол ~90°)
        elif dx < dy * 0.2:
            x_avg = (x1 + x2) // 2
            if margin <= x_avg <= w - margin:
                v_coords.add(x_avg)

    return sorted(h_coords), sorted(v_coords)


def _cluster_coords(coords: list[int], threshold: int = 10) -> list[int]:
    """Кластеризует близкие координаты: если разница ≤ threshold — объединяем."""
    if not coords:
        return []
    clusters: list[int] = []
    current = coords[0]
    count = 1
    for c in coords[1:]:
        if c - current <= threshold:
            current = (current * count + c) // (count + 1)
            count += 1
        else:
            clusters.append(current)
            current = c
            count = 1
    clusters.append(current)
    return clusters


def detect_table_grid(
    img: np.ndarray
) -> tuple[list[int], list[int]] | None:
    """
    Детектирует сетку таблицы на изображении.
    Возвращает (rows_y, cols_x) — координаты линий строк и столбцов,
    или None, если таблица не найдена.
    """
    binary = _preprocess_for_lines(img)
    h_coords, v_coords = _find_lines(binary)

    # Кластеризация
    rows = _cluster_coords(h_coords, threshold=12)
    cols = _cluster_coords(v_coords, threshold=12)

    # Для таблицы нужно как минимум 2 строки и 2 столбца
    if len(rows) < 2 or len(cols) < 2:
        return None

    return rows, cols


def cells_to_markdown(
    img: np.ndarray,
    rows: list[int],
    cols: list[int],
    reader: Any
) -> str:
    """
    Разбивает изображение на ячейки по сетке (rows × cols),
    распознаёт текст в каждой ячейке и собирает Markdown-таблицу.

    Первая строка ячеек считается заголовком таблицы.
    """
    table: list[list[str]] = []
    for ri in range(len(rows) - 1):
        y1 = max(0, rows[ri])
        y2 = min(img.shape[0], rows[ri + 1])
        # Пропускаем слишком узкие строки
        if y2 - y1 < 8:
            continue

        row_cells: list[str] = []
        for ci in range(len(cols) - 1):
            x1 = max(0, cols[ci])
            x2 = min(img.shape[1], cols[ci + 1])
            # Пропускаем слишком узкие столбцы
            if x2 - x1 < 8:
                continue

            cell_img = img[y1:y2, x1:x2]
            results = reader.readtext(cell_img, detail=0, paragraph=True)
            text = ' '.join(results).strip()
            row_cells.append(text)

        # Пропускаем пустые строки
        if not any(c.strip() for c in row_cells):
            continue
        # Добиваем до числа столбцов, если не хватает
        while len(row_cells) < len(cols) - 1:
            row_cells.append('')
        table.append(row_cells)

    if not table:
        return ''

    # Сборка Markdown
    md_lines: list[str] = []

    # Заголовок
    if table:
        md_lines.append('| ' + ' | '.join(table[0]) + ' |')

    # Разделитель заголовка
    sep = '| ' + ' | '.join(['---'] * len(table[0])) + ' |'
    md_lines.append(sep)

    # Тело таблицы
    for row in table[1:]:
        # Выравниваем длину строки по числу столбцов заголовка
        while len(row) < len(table[0]):
            row.append('')
        md_lines.append('| ' + ' | '.join(row[:len(table[0])]) + ' |')

    return '\n'.join(md_lines)


def extract_table_from_image(data: bytes) -> tuple[str, str]:
    """
    Изображение → детекция таблицы → OCR по ячейкам → Markdown.
    Если таблица не найдена — fallback на обычный EasyOCR.
    """
    import numpy as np
    from PIL import Image
    reader = get_easyocr()
    img = Image.open(io.BytesIO(data))
    if img.mode not in ('RGB', 'L'):
        img = img.convert('RGB')
    arr = np.array(img)

    grid = detect_table_grid(arr)
    if grid is None:
        # Таблица не найдена — обычный OCR
        log.info('Table not detected, falling back to plain OCR')
        result = reader.readtext(arr, detail=0, paragraph=True)
        text = '\n'.join(result).strip()
        return text, 'easyocr'

    rows, cols = grid
    log.info('Table detected: %d rows × %d cols', len(rows) - 1, len(cols) - 1)
    markdown = cells_to_markdown(arr, rows, cols, reader)
    if not markdown:
        # Пустая таблица — тоже fallback
        result = reader.readtext(arr, detail=0, paragraph=True)
        text = '\n'.join(result).strip()
        return text, 'easyocr'

    return markdown, 'easyocr+table'


def extract_table_from_pdf(data: bytes) -> tuple[str, str]:
    """
    PDF: сначала пробуем прямое извлечение (PyMuPDF).
    Если текста мало — рендерим страницы и детектируем таблицы через OCR.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(stream=data, filetype='pdf')
    try:
        if doc.page_count > PDF_MAX_PAGES:
            log.warning('PDF has %d pages, limiting to %d', doc.page_count, PDF_MAX_PAGES)

        # Шаг 1: прямое извлечение
        parts = []
        for i, page in enumerate(doc):
            if i >= PDF_MAX_PAGES:
                break
            parts.append(page.get_text('text'))
        text = '\n'.join(parts).strip()

        if len(text) >= PDF_MIN_TEXT:
            return text, 'pymupdf'

        # Шаг 2: рендерим страницы и детектируем таблицы
        log.info('PDF text too short (%d chars), rendering for table detection', len(text))
        import numpy as np
        from PIL import Image
        reader = get_easyocr()

        all_pages: list[str] = []
        for i, page in enumerate(doc):
            if i >= PDF_MAX_PAGES:
                break
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_bytes = pix.tobytes('png')
            img = Image.open(io.BytesIO(img_bytes))
            arr = np.array(img)

            grid = detect_table_grid(arr)
            if grid is not None:
                rows, cols = grid
                log.info('Page %d: table detected (%d rows × %d cols)',
                         i + 1, len(rows) - 1, len(cols) - 1)
                md = cells_to_markdown(arr, rows, cols, reader)
                if md:
                    all_pages.append(md)
                    continue

            # Fallback: обычный OCR для страницы
            result = reader.readtext(arr, detail=0, paragraph=True)
            all_pages.append('\n'.join(result))

        full = '\n\n---\n\n'.join(all_pages).strip()
        return full, 'pymupdf+easyocr+table'
    finally:
        doc.close()


# --- Извлечение текста по типам ---

def extract_pdf(data: bytes) -> tuple[str, str]:
    """
    PDF: сначала PyMuPDF (прямое извлечение). Если получено мало текста —
    рендерим страницы как картинки и пропускаем через EasyOCR.
    Возвращает (text, source).
    """
    import fitz  # PyMuPDF

    doc = fitz.open(stream=data, filetype='pdf')
    try:
        if doc.page_count > PDF_MAX_PAGES:
            log.warning('PDF has %d pages, limiting to %d', doc.page_count, PDF_MAX_PAGES)

        # Шаг 1: прямое извлечение через PyMuPDF
        parts = []
        for i, page in enumerate(doc):
            if i >= PDF_MAX_PAGES:
                break
            parts.append(page.get_text('text'))
        text = '\n'.join(parts).strip()

        if len(text) >= PDF_MIN_TEXT:
            return text, 'pymupdf'

        # Шаг 2: PDF без текстового слоя — рендерим страницы и OCR-им
        log.info('PDF text too short (%d chars), falling back to EasyOCR', len(text))
        reader = get_easyocr()
        ocr_parts = []
        for i, page in enumerate(doc):
            if i >= PDF_MAX_PAGES:
                break
            # 2x — увеличенное разрешение для лучшего распознавания
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img_bytes = pix.tobytes('png')
            # EasyOCR принимает numpy-array или путь к файлу
            import numpy as np
            from PIL import Image
            img = Image.open(io.BytesIO(img_bytes))
            arr = np.array(img)
            result = reader.readtext(arr, detail=0, paragraph=True)
            ocr_parts.append('\n'.join(result))
        full = '\n'.join(ocr_parts).strip()
        return full, 'pymupdf+easyocr'
    finally:
        doc.close()


def extract_docx(data: bytes) -> tuple[str, str]:
    """DOCX: через python-docx (читаем параграфы + таблицы)."""
    import docx
    f = io.BytesIO(data)
    doc = docx.Document(f)
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    text = '\n'.join(parts).strip()
    return text, 'docx'


def extract_image(data: bytes) -> tuple[str, str]:
    """Картинки → EasyOCR."""
    import numpy as np
    from PIL import Image
    reader = get_easyocr()
    img = Image.open(io.BytesIO(data))
    # Гарантируем RGB (EasyOCR не любит RGBA)
    if img.mode not in ('RGB', 'L'):
        img = img.convert('RGB')
    arr = np.array(img)
    result = reader.readtext(arr, detail=0, paragraph=True)
    text = '\n'.join(result).strip()
    return text, 'easyocr'


def extract_plain(data: bytes) -> tuple[str, str]:
    """TXT/MD/CSV: пробуем UTF-8, потом cp1251 (распространено в офисных файлах)."""
    for enc in ('utf-8', 'utf-8-sig', 'cp1251', 'latin-1'):
        try:
            return data.decode(enc).strip(), 'plain'
        except UnicodeDecodeError:
            continue
    raise HTTPException(415, 'Не удалось декодировать текстовый файл')


# --- FastAPI ---

app = FastAPI(title='GigaChat OCR Server')


@app.get('/status')
def status():
    """Лёгкий health-check. НЕ инициализирует EasyOCR."""
    return {
        'status': 'ok',
        'easyocr_ready': _easyocr_reader is not None,
        'easyocr_error': _easyocr_init_error,
        'device': _detect_device(),
        'langs': LANGS,
        'port': PORT,
    }


IMAGE_EXTS = {'png', 'jpg', 'jpeg', 'tiff', 'tif', 'bmp', 'webp'}
PLAIN_EXTS = {'txt', 'md', 'csv', 'log'}


@app.post('/extract')
async def extract(file: UploadFile = File(...)):
    """
    Основной эндпоинт. Принимает любой документ, отдаёт извлечённый текст.
    Совместим с тем, что ожидают workflow'ы organization-appeal, document-loader и др.
    """
    if not file.filename:
        raise HTTPException(400, 'Файл без имени')
    # M1: читаем с лимитом размера — защита от OOM при прямом POST в обход n8n
    # (сервер слушает 0.0.0.0 в LAN). Чанковое чтение: не тянем весь файл в
    # память целиком, если он превышает лимит.
    max_bytes = MAX_UPLOAD_MB * 1024 * 1024
    buf = bytearray()
    while True:
        chunk = await file.read(1024 * 1024)
        if not chunk:
            break
        buf.extend(chunk)
        if len(buf) > max_bytes:
            raise HTTPException(413, f'Файл больше {MAX_UPLOAD_MB} МБ — отклонён.')
    data = bytes(buf)
    if not data:
        raise HTTPException(400, 'Пустой файл')

    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    log.info('Received: %s (%d bytes, ext=%s)', file.filename, len(data), ext)

    try:
        if ext == 'pdf':
            text, source = extract_pdf(data)
        elif ext == 'docx':
            text, source = extract_docx(data)
        elif ext in IMAGE_EXTS:
            text, source = extract_image(data)
        elif ext in PLAIN_EXTS:
            text, source = extract_plain(data)
        else:
            raise HTTPException(415, f'Неподдерживаемый формат: .{ext}. Поддерживаются: pdf, docx, txt, png, jpg, tiff, bmp, webp.')
    except HTTPException:
        raise
    except RuntimeError as e:
        # EasyOCR не загрузился — отдадим понятную ошибку
        raise HTTPException(503, str(e))
    except Exception as e:
        log.exception('Extract error')
        raise HTTPException(500, f'Ошибка извлечения: {e}')

    log.info('Extracted %d chars via %s', len(text), source)
    return {
        'text': text,
        'source': source,
        'filename': file.filename,
        'chars': len(text),
    }


@app.post('/extract-table')
async def extract_table(file: UploadFile = File(...)):
    """
    Эндпоинт детекции таблиц. Принимает PDF или изображение,
    определяет структуру таблицы и возвращает текст в формате Markdown.
    Для изображений без таблиц — fallback на обычный OCR (плоский текст).
    """
    if not file.filename:
        raise HTTPException(400, 'Файл без имени')
    max_bytes = MAX_UPLOAD_MB * 1024 * 1024
    buf = bytearray()
    while True:
        chunk = await file.read(1024 * 1024)
        if not chunk:
            break
        buf.extend(chunk)
        if len(buf) > max_bytes:
            raise HTTPException(413, f'Файл больше {MAX_UPLOAD_MB} МБ — отклонён.')
    data = bytes(buf)
    if not data:
        raise HTTPException(400, 'Пустой файл')

    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    log.info('extract-table received: %s (%d bytes, ext=%s)', file.filename, len(data), ext)

    try:
        if ext == 'pdf':
            text, source = extract_table_from_pdf(data)
        elif ext in IMAGE_EXTS:
            text, source = extract_table_from_image(data)
        else:
            raise HTTPException(415,
                f'Неподдерживаемый формат: .{ext}. '
                f'extract-table поддерживает: pdf, png, jpg, tiff, bmp, webp.')
    except HTTPException:
        raise
    except RuntimeError as e:
        raise HTTPException(503, str(e))
    except Exception as e:
        log.exception('extract-table error')
        raise HTTPException(500, f'Ошибка извлечения таблицы: {e}')

    log.info('extract-table extracted %d chars via %s', len(text), source)
    return {
        'text': text,
        'source': source,
        'filename': file.filename,
        'chars': len(text),
    }


if __name__ == '__main__':
    log.info('Starting OCR server on http://%s:%d', HOST, PORT)
    log.info('Langs: %s | PDF min text: %d | PDF max pages: %d',
             LANGS, PDF_MIN_TEXT, PDF_MAX_PAGES)
    uvicorn.run(app, host=HOST, port=PORT, log_level='info')

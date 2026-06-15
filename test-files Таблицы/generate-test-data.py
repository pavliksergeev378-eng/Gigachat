"""
Генератор тестовых данных для table-merger.

Создаёт 5 xlsx и 5 docx файлов в текущей папке.

Все файлы — таблица «сотрудники компании», но с РАЗНЫМ набором столбцов
и РАЗНЫМ написанием одинаковых по смыслу заголовков. Это нагрузочный тест
для нормализатора заголовков и логики union столбцов:

  - Нормализация склеит:
      ФИО / Ф.И.О. / фио / Ф И О → один столбец
      Email / E-mail               → один столбец

  - Нормализация НЕ склеит (это разные слова — союз сохранит как разные):
      Дата приёма / Дата найма / Дата трудоустройства  → 3 столбца
      Стаж работы (лет) / Стаж, лет                    → 2 столбца

  - Union сохранит уникальные столбцы из отдельных файлов:
      Телефон, Email, Возраст, Образование, Бонус %, Стаж — каждый есть
      не во всех файлах, в строках без них итог будет пустой.

В файлах часть значений намеренно пустые (~10% ячеек), чтобы протестировать
обработку пропусков.

Запуск:
    cd C:\\Users\\Lenovo\\Desktop\\GigaChat\\test-files
    python generate-test-data.py

Требуются: openpyxl, python-docx
    pip install openpyxl python-docx
"""

import random
from datetime import date, timedelta
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document


# Фиксируем seed чтобы при повторном запуске получались те же файлы.
random.seed(42)


NAMES_M = [
    'Иванов Иван Иванович',
    'Петров Пётр Петрович',
    'Сидоров Алексей Николаевич',
    'Кузнецов Дмитрий Сергеевич',
    'Смирнов Владимир Андреевич',
    'Соколов Михаил Викторович',
    'Попов Андрей Дмитриевич',
    'Лебедев Сергей Александрович',
    'Козлов Артём Олегович',
    'Новиков Илья Романович',
    'Морозов Григорий Павлович',
    'Волков Антон Игоревич',
]
NAMES_F = [
    'Иванова Анна Сергеевна',
    'Петрова Мария Викторовна',
    'Сидорова Ольга Дмитриевна',
    'Кузнецова Екатерина Игоревна',
    'Смирнова Татьяна Николаевна',
    'Соколова Наталья Андреевна',
    'Попова Елена Александровна',
    'Лебедева Светлана Петровна',
    'Козлова Юлия Романовна',
    'Новикова Дарья Викторовна',
    'Морозова Анастасия Олеговна',
    'Волкова Полина Игоревна',
]

POSITIONS = [
    'Менеджер', 'Старший менеджер', 'Аналитик', 'Старший аналитик',
    'Разработчик', 'Старший разработчик', 'Тимлид', 'Дизайнер',
    'UX-исследователь', 'Бухгалтер', 'Финансовый аналитик',
    'Тестировщик', 'Специалист по продажам', 'HR-специалист',
    'Руководитель отдела',
]

DEPARTMENTS = [
    'Маркетинг', 'Разработка', 'Финансы',
    'Продажи', 'HR', 'Поддержка', 'Аналитика', 'Дизайн',
]

EDUCATION = [
    'Высшее техническое',
    'Высшее экономическое',
    'Высшее гуманитарное',
    'Магистратура',
    'Бакалавр',
    'Среднее специальное',
    'Незаконченное высшее',
]


# Транслитерация фамилии для email. Простая ASCII-table.
_TRANSLIT = {
    'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd', 'е': 'e', 'ё': 'e',
    'ж': 'zh', 'з': 'z', 'и': 'i', 'й': 'y', 'к': 'k', 'л': 'l', 'м': 'm',
    'н': 'n', 'о': 'o', 'п': 'p', 'р': 'r', 'с': 's', 'т': 't', 'у': 'u',
    'ф': 'f', 'х': 'h', 'ц': 'ts', 'ч': 'ch', 'ш': 'sh', 'щ': 'sch',
    'ъ': '', 'ы': 'y', 'ь': '', 'э': 'e', 'ю': 'yu', 'я': 'ya',
}


def translit(s: str) -> str:
    return ''.join(_TRANSLIT.get(c.lower(), c.lower()) for c in s if c.isalpha())


# Описание файлов: (имя, заголовки).
# Намеренно ОЧЕНЬ разные структуры — от минимума (5 столбцов) до полного
# набора (9 столбцов). Разные написания одного и того же столбца.
FILE_VARIATIONS = [
    # Минимальный — только базовые поля
    ('маркетинг',  [
        'ФИО', 'Должность', 'Отдел', 'Зарплата', 'Дата приёма',
    ]),
    # Базовый + контакты + стаж, ФИО через точки
    ('разработка', [
        'Ф.И.О.', 'Должность', 'Отдел', 'Зарплата', 'Дата приёма',
        'Email', 'Телефон', 'Стаж работы (лет)',
    ]),
    # Все заголовки в нижнем регистре + образование + ИНАЯ КОЛОНКА «Дата найма»
    ('финансы',    [
        'фио', 'должность', 'отдел', 'зарплата', 'дата найма',
        'образование',
    ]),
    # Самый широкий — 9 столбцов, с возрастом и бонусом
    ('продажи',    [
        'ФИО', 'Должность', 'Отдел', 'Зарплата', 'Дата приёма',
        'Телефон', 'Email', 'Возраст', 'Бонус, %',
    ]),
    # ФИО через пробелы + ЕЩЁ ОДНА вариация даты + email с дефисом + стаж по-другому
    ('hr',         [
        'Ф И О', 'Должность', 'Отдел', 'Зарплата', 'Дата трудоустройства',
        'E-mail', 'Стаж, лет',
    ]),
]


def normalize_for_match(h: str) -> str:
    """Тот же нормализатор что в браузерном table-merger.
    Используем здесь чтобы понимать какой смысл несёт каждый заголовок."""
    import re
    return re.sub(r"[\s._\-(){}\[\]/\\:;,!?\"' ]+", '', str(h).strip().lower())


def gen_value_for_header(header: str, name_for_email: str = '') -> object:
    """По имени заголовка генерируем подходящее значение.
    name_for_email — фамилия чтобы сделать осмысленный email."""
    key = normalize_for_match(header)

    # ~10% значений делаем пустыми (NULL) — для теста пропусков.
    # Кроме обязательных полей (ФИО, Должность, Отдел).
    is_optional = not any(k in key for k in ('фио', 'должность', 'отдел'))
    if is_optional and random.random() < 0.1:
        return ''

    if 'фио' in key:
        return random.choice(NAMES_M + NAMES_F)
    if 'должность' in key:
        return random.choice(POSITIONS)
    if 'отдел' in key:
        return random.choice(DEPARTMENTS)
    if 'зарплата' in key:
        # Иногда круглые тысячи, иногда копейки
        if random.random() < 0.3:
            return round(random.uniform(50_000, 300_000), 2)
        return random.randint(50_000, 300_000)
    # Все три варианта даты («дата приёма», «дата найма», «дата трудоустройства»)
    # генерируем одинаково — это просто разные слова для одного смысла.
    if 'дата' in key:
        d = date(2018, 1, 1) + timedelta(days=random.randint(0, 365 * 7))
        return d.strftime('%d.%m.%Y')
    if 'телефон' in key:
        return '+7' + ''.join(str(random.randint(0, 9)) for _ in range(10))
    if 'email' in key or 'mail' in key:
        # фамилия@company.ru — translit фамилии
        surname = (name_for_email or '').split()[0] if name_for_email else 'user'
        return f'{translit(surname)}@gigachat-corp.ru'
    if 'возраст' in key:
        return random.randint(22, 62)
    if 'образование' in key:
        return random.choice(EDUCATION)
    if 'бонус' in key:
        return random.randint(0, 30)
    if 'стаж' in key:
        return random.randint(0, 18)
    return '?'


def gen_row(headers: list) -> list:
    """Генерируем одну строку, передавая имя дальше в gen для email."""
    # Сначала ФИО (для использования в email)
    name = ''
    for h in headers:
        if 'фио' in normalize_for_match(h):
            name = random.choice(NAMES_M + NAMES_F)
            break

    row = []
    used_name = False
    for h in headers:
        if 'фио' in normalize_for_match(h):
            if not used_name and name:
                row.append(name)
                used_name = True
            else:
                row.append(random.choice(NAMES_M + NAMES_F))
        else:
            row.append(gen_value_for_header(h, name))
    return row


# ---------- XLSX ----------

def make_xlsx(filepath: Path, headers: list, n_rows: int) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = 'Сотрудники'

    header_font = Font(bold=True, color='FFFFFFFF')
    header_fill = PatternFill('solid', fgColor='FF4F46E5')
    header_align = Alignment(horizontal='center', vertical='center', wrap_text=True)

    for col_idx, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align

    for r in range(n_rows):
        row_data = gen_row(headers)
        for col_idx, val in enumerate(row_data, start=1):
            ws.cell(row=r + 2, column=col_idx, value=val)

    # Авто-ширина (грубая)
    for col_idx, h in enumerate(headers, start=1):
        max_len = max(len(str(h)), 12)
        for r in range(n_rows):
            v = ws.cell(row=r + 2, column=col_idx).value
            if v is not None and len(str(v)) > max_len:
                max_len = len(str(v))
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = min(max_len + 2, 35)

    ws.freeze_panes = 'A2'
    wb.save(filepath)


# ---------- DOCX ----------

def make_docx(filepath: Path, headers: list, n_rows: int, dept_label: str) -> None:
    doc = Document()
    doc.add_heading(f'Отчёт по отделу: {dept_label}', level=1)
    doc.add_paragraph(
        'Источник: HR-отдел. Период: 2024 год. '
        'Документ сгенерирован автоматически для тестирования table-merger.'
    )
    doc.add_paragraph()

    table = doc.add_table(rows=1 + n_rows, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for r in range(n_rows):
        row_data = gen_row(headers)
        row_cells = table.rows[r + 1].cells
        for col_idx, val in enumerate(row_data):
            row_cells[col_idx].text = '' if val is None or val == '' else str(val)

    doc.add_paragraph()
    doc.add_paragraph(
        'Этот текст идёт ПОСЛЕ таблицы — он также должен быть проигнорирован.'
    )
    doc.save(filepath)


# ---------- main ----------

def make_xlsx_multi_sheet(filepath: Path, sheets_spec: list) -> None:
    """Excel с несколькими листами. sheets_spec: [(sheet_name, headers, n_rows), ...]"""
    wb = Workbook()
    # Удалим default-лист, добавим свои
    wb.remove(wb.active)
    header_font = Font(bold=True, color='FFFFFFFF')
    header_fill = PatternFill('solid', fgColor='FF4F46E5')
    header_align = Alignment(horizontal='center', vertical='center', wrap_text=True)

    for sheet_name, headers, n_rows in sheets_spec:
        ws = wb.create_sheet(title=sheet_name)
        for col_idx, h in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=col_idx, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
        for r in range(n_rows):
            row_data = gen_row(headers)
            for col_idx, val in enumerate(row_data, start=1):
                ws.cell(row=r + 2, column=col_idx, value=val)
        ws.freeze_panes = 'A2'
        for col_idx, h in enumerate(headers, start=1):
            max_len = max(len(str(h)), 12)
            for r in range(n_rows):
                v = ws.cell(row=r + 2, column=col_idx).value
                if v is not None and len(str(v)) > max_len:
                    max_len = len(str(v))
            ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = min(max_len + 2, 35)
    wb.save(filepath)


def make_docx_multi_table(filepath: Path, tables_spec: list, doc_title: str) -> None:
    """Word с несколькими таблицами. tables_spec: [(label, headers, n_rows), ...]"""
    doc = Document()
    doc.add_heading(doc_title, level=1)
    doc.add_paragraph('Документ с несколькими таблицами — для теста multi-table в Word-мерджере.')

    for label, headers, n_rows in tables_spec:
        doc.add_heading(label, level=2)
        table = doc.add_table(rows=1 + n_rows, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for r in range(n_rows):
            row_data = gen_row(headers)
            row_cells = table.rows[r + 1].cells
            for col_idx, val in enumerate(row_data):
                row_cells[col_idx].text = '' if val is None or val == '' else str(val)
        doc.add_paragraph()
    doc.save(filepath)


def make_xlsx_with_duplicates(filepath: Path, headers: list, n_rows: int) -> None:
    """Excel со строками, которые повторятся при объединении с другими файлами.
    10 «общих» строк с фиксированными значениями + ВНУТРЕННИЕ дубли (повторяем
    каждую общую строку дважды). Итого при загрузке файла → 20 строк дублей,
    при объединении нескольких файлов — кратно больше. Для теста dedup."""
    common_rows = [
        ['Иванов Иван Иванович',       'Менеджер',           'Продажи',     75000, '15.03.2020'],
        ['Петрова Мария Викторовна',   'Аналитик',           'Аналитика',   95000, '01.06.2019'],
        ['Сидоров Алексей Николаевич', 'Тимлид',             'Разработка', 180000, '12.09.2018'],
        ['Кузнецов Дмитрий Сергеевич', 'Дизайнер',           'Дизайн',     110000, '04.02.2021'],
        ['Смирнов Владимир Андреевич', 'Бухгалтер',          'Финансы',     85000, '20.11.2019'],
        ['Соколов Михаил Викторович',  'Тестировщик',        'Разработка',  90000, '07.07.2022'],
        ['Попов Андрей Дмитриевич',    'HR-специалист',      'HR',          78000, '15.01.2020'],
        ['Лебедев Сергей Александрович','Старший разработчик','Разработка', 220000, '03.05.2017'],
        ['Иванова Анна Сергеевна',     'UX-исследователь',   'Дизайн',     130000, '11.08.2021'],
        ['Козлова Юлия Романовна',     'Финансовый аналитик','Финансы',    140000, '22.04.2018'],
    ]
    # Каждая common-строка появляется ДВАЖДЫ для внутренних дублей
    common_rows = common_rows + common_rows
    wb = Workbook()
    ws = wb.active
    ws.title = 'Сотрудники'

    header_font = Font(bold=True, color='FFFFFFFF')
    header_fill = PatternFill('solid', fgColor='FF4F46E5')

    for col_idx, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx, value=h)
        cell.font = header_font
        cell.fill = header_fill

    # Сначала «общие» (для дублирования с другими файлами)
    for r, row_data in enumerate(common_rows):
        for col_idx, val in enumerate(row_data, start=1):
            ws.cell(row=r + 2, column=col_idx, value=val)

    # Потом обычные сгенерированные
    for r in range(n_rows):
        row_data = gen_row(headers)
        for col_idx, val in enumerate(row_data, start=1):
            ws.cell(row=r + 2 + len(common_rows), column=col_idx, value=val)

    ws.freeze_panes = 'A2'
    wb.save(filepath)


def make_docx_with_duplicates(filepath: Path, headers: list, n_rows: int, dept_label: str) -> None:
    """Word аналог make_xlsx_with_duplicates — для теста dedup.
    10 общих строк × 2 повтора = 20 внутренних дублей."""
    common_rows = [
        ['Иванов Иван Иванович',       'Менеджер',           'Продажи',     '75000',  '15.03.2020'],
        ['Петрова Мария Викторовна',   'Аналитик',           'Аналитика',   '95000',  '01.06.2019'],
        ['Сидоров Алексей Николаевич', 'Тимлид',             'Разработка',  '180000', '12.09.2018'],
        ['Кузнецов Дмитрий Сергеевич', 'Дизайнер',           'Дизайн',      '110000', '04.02.2021'],
        ['Смирнов Владимир Андреевич', 'Бухгалтер',          'Финансы',     '85000',  '20.11.2019'],
        ['Соколов Михаил Викторович',  'Тестировщик',        'Разработка',  '90000',  '07.07.2022'],
        ['Попов Андрей Дмитриевич',    'HR-специалист',      'HR',          '78000',  '15.01.2020'],
        ['Лебедев Сергей Александрович','Старший разработчик','Разработка', '220000', '03.05.2017'],
        ['Иванова Анна Сергеевна',     'UX-исследователь',   'Дизайн',      '130000', '11.08.2021'],
        ['Козлова Юлия Романовна',     'Финансовый аналитик','Финансы',     '140000', '22.04.2018'],
    ]
    common_rows = common_rows + common_rows  # каждую — дважды для внутренних дублей
    doc = Document()
    doc.add_heading(f'Отчёт: {dept_label} (с дублями)', level=1)
    doc.add_paragraph('Этот файл содержит строки, которые есть и в filiale.docx — для теста удаления дубликатов.')

    table = doc.add_table(rows=1 + n_rows + len(common_rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True

    for r, row_data in enumerate(common_rows):
        row_cells = table.rows[r + 1].cells
        for col_idx, val in enumerate(row_data):
            row_cells[col_idx].text = str(val)

    for r in range(n_rows):
        row_data = gen_row(headers)
        row_cells = table.rows[r + 1 + len(common_rows)].cells
        for col_idx, val in enumerate(row_data):
            row_cells[col_idx].text = '' if val is None or val == '' else str(val)

    doc.save(filepath)


def main() -> None:
    out_dir = Path(__file__).parent
    print(f'Генерирую файлы в: {out_dir}')
    print()

    summary = []
    for idx, (name, headers) in enumerate(FILE_VARIATIONS, start=1):
        n_rows_x = random.randint(5, 15)
        xlsx_path = out_dir / f'{name}.xlsx'
        make_xlsx(xlsx_path, headers, n_rows_x)

        n_rows_d = random.randint(5, 15)
        docx_path = out_dir / f'{name}.docx'
        make_docx(docx_path, headers, n_rows_d, name.capitalize())

        summary.append((name, headers, n_rows_x, n_rows_d))
        print(f'  [{idx}/{len(FILE_VARIATIONS)}] {name}: '
              f'{len(headers)} колонок | xlsx={n_rows_x} строк | docx={n_rows_d} строк')

    # ============ Дополнительные файлы под новые фичи ============

    print()
    print('Дополнительные файлы (для новых фич мерджеров):')

    # 1. Excel с 3 листами (multi-sheet)
    multi_xlsx = out_dir / 'multi-sheet.xlsx'
    make_xlsx_multi_sheet(multi_xlsx, [
        ('Q1', ['ФИО', 'Должность', 'Отдел', 'Зарплата', 'Дата приёма'], 6),
        ('Q2', ['ФИО', 'Должность', 'Отдел', 'Зарплата', 'Дата приёма', 'Бонус, %'], 7),
        ('Q3', ['Ф.И.О.', 'должность', 'отдел', 'зарплата', 'Email'], 5),
    ])
    print(f'  multi-sheet.xlsx — 3 листа (Q1/Q2/Q3) с разными заголовками')

    # 2. Word с 3 таблицами (multi-table)
    multi_docx = out_dir / 'multi-table.docx'
    make_docx_multi_table(multi_docx, [
        ('Январь', ['ФИО', 'Должность', 'Отдел', 'Зарплата', 'Дата приёма'], 6),
        ('Февраль', ['ФИО', 'Должность', 'Отдел', 'Зарплата', 'Дата приёма', 'Бонус, %'], 7),
        ('Март', ['Ф.И.О.', 'должность', 'отдел', 'зарплата', 'Email'], 5),
    ], 'Квартальный отчёт по сотрудникам')
    print(f'  multi-table.docx — 3 таблицы (Январь/Февраль/Март) с разными заголовками')

    # 3. Файлы с заранее заданными «общими» строками для dedup-теста
    dup_xlsx = out_dir / 'дубли.xlsx'
    make_xlsx_with_duplicates(dup_xlsx, [
        'ФИО', 'Должность', 'Отдел', 'Зарплата', 'Дата приёма',
    ], 5)
    dup_docx = out_dir / 'дубли.docx'
    make_docx_with_duplicates(dup_docx, [
        'ФИО', 'Должность', 'Отдел', 'Зарплата', 'Дата приёма',
    ], 5, 'Филиал')
    print(f'  дубли.xlsx / .docx — содержат 3 общие строки (Иванов/Петрова/Сидоров)')
    print(f'                       при merge с hr.xlsx/.docx (нет общих) или с самим собой')
    print(f'                       произведут дубликаты — для теста галки «Удалить дубликаты»')

    print()
    print('Готово.')
    print()
    print('Сводка по заголовкам (видно вариативность для теста):')
    for name, headers, _, _ in summary:
        print(f'  {name}: {headers}')

    print()
    print('=== Сценарии тестирования ===')
    print()
    print('[1] Базовый union (объединить все 5 одного формата):')
    print('     - ФИО (4 формы) → 1 колонка; Дата (3 формы) → 3 колонки; и т.д.')
    print()
    print('[2] Multi-sheet Excel: загрузить multi-sheet.xlsx — badge «3 лист.»,')
    print('     объединение даст единую таблицу со union колонок всех 3 листов.')
    print()
    print('[3] Multi-table Word: аналогично с multi-table.docx — badge «3 табл.».')
    print()
    print('[4] Dedup: загрузить дубли.xlsx 2 раза (или дубли + что-то ещё) →')
    print('     отметить «Удалить дубликаты» → 3 «общие» строки схлопнутся в 1 каждая.')
    print()
    print('[5] Sort: после merge выбрать сортировку по «Зарплата» (числовая)')
    print('     или «Дата приёма» (строковая) — увидеть упорядоченный результат.')
    print()
    print('[6] Ignore columns: после merge кликнуть на pill'
          ' «Бонус, %» или «Образование» — они выпадут из превью и результата.')
    print()
    print('[7] Drag-reorder: в списке файлов перетащить за ⠿ — порядок merge меняется.')


if __name__ == '__main__':
    main()

# -*- coding: utf-8 -*-
# Генератор тестовых файлов для Agents/file-merger.html.
# Создаёт 5 .docx (текст / таблица / mixed / multiple tables / long text)
# и 5 .xlsx (разной структуры) в этой папке.
# Использует только stdlib (zipfile + xml-строки), без внешних библиотек.
#
# Запуск:  python "Тест. Объединение файлов/_generate.py"

import os
import sys
import zipfile
from pathlib import Path

TD = Path(__file__).parent.resolve()


# ============ Утилиты ============

def esc(s):
    return str(s).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


# ============ DOCX ============

def make_docx(filename, body_xml):
    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        '</Types>'
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
        '</Relationships>'
    )
    document = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:body>'
        + body_xml +
        '<w:sectPr><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1134" w:right="850" w:bottom="1134" w:left="1701" w:header="708" w:footer="708" w:gutter="0"/></w:sectPr>'
        '</w:body></w:document>'
    )
    path = TD / filename
    with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('[Content_Types].xml', content_types)
        zf.writestr('_rels/.rels', rels)
        zf.writestr('word/document.xml', document)
    print('  + ' + filename)


def P(text):
    return '<w:p><w:r><w:t xml:space="preserve">' + esc(text) + '</w:t></w:r></w:p>'


def TBL(rows):
    cols = max(len(r) for r in rows)
    cell_w = 9000 // cols
    grid = ''.join('<w:gridCol w:w="' + str(cell_w) + '"/>' for _ in range(cols))
    tr_xml = ''
    for row in rows:
        tc_xml = ''
        for i in range(cols):
            cell = row[i] if i < len(row) else ''
            tc_xml += (
                '<w:tc><w:tcPr><w:tcW w:w="' + str(cell_w) + '" w:type="dxa"/></w:tcPr>'
                '<w:p><w:r><w:t xml:space="preserve">' + esc(cell) + '</w:t></w:r></w:p>'
                '</w:tc>'
            )
        tr_xml += '<w:tr>' + tc_xml + '</w:tr>'
    borders = ''.join(
        '<w:' + side + ' w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV')
    )
    return (
        '<w:tbl>'
        '<w:tblPr><w:tblW w:w="9000" w:type="dxa"/><w:tblBorders>' + borders + '</w:tblBorders></w:tblPr>'
        '<w:tblGrid>' + grid + '</w:tblGrid>'
        + tr_xml +
        '</w:tbl>'
    )


# ============ XLSX ============

def col_letter(n):
    s = ''
    while n >= 0:
        s = chr(ord('A') + n % 26) + s
        n = n // 26 - 1
    return s


def make_xlsx(filename, sheet_name, rows):
    strings = {}
    string_list = []
    for row in rows:
        for cell in row:
            if isinstance(cell, str):
                if cell not in strings:
                    strings[cell] = len(string_list)
                    string_list.append(cell)

    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
        '<Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
        '<Override PartName="/xl/sharedStrings.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sharedStrings+xml"/>'
        '</Types>'
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>'
        '</Relationships>'
    )
    workbook = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<sheets><sheet name="' + esc(sheet_name) + '" sheetId="1" r:id="rId1"/></sheets>'
        '</workbook>'
    )
    wb_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>'
        '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/sharedStrings" Target="sharedStrings.xml"/>'
        '</Relationships>'
    )
    rows_xml = ''
    for r_idx, row in enumerate(rows):
        cells_xml = ''
        for c_idx, cell in enumerate(row):
            ref = col_letter(c_idx) + str(r_idx + 1)
            if isinstance(cell, str):
                cells_xml += '<c r="' + ref + '" t="s"><v>' + str(strings[cell]) + '</v></c>'
            else:
                cells_xml += '<c r="' + ref + '"><v>' + str(cell) + '</v></c>'
        rows_xml += '<row r="' + str(r_idx + 1) + '">' + cells_xml + '</row>'
    sheet1 = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        '<sheetData>' + rows_xml + '</sheetData>'
        '</worksheet>'
    )
    total = sum(1 for r in rows for c in r if isinstance(c, str))
    si_xml = ''.join('<si><t xml:space="preserve">' + esc(s) + '</t></si>' for s in string_list)
    sst = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<sst xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" count="' + str(total) + '" uniqueCount="' + str(len(string_list)) + '">'
        + si_xml +
        '</sst>'
    )
    path = TD / filename
    with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('[Content_Types].xml', content_types)
        zf.writestr('_rels/.rels', rels)
        zf.writestr('xl/workbook.xml', workbook)
        zf.writestr('xl/_rels/workbook.xml.rels', wb_rels)
        zf.writestr('xl/worksheets/sheet1.xml', sheet1)
        zf.writestr('xl/sharedStrings.xml', sst)
    print('  + ' + filename)


# ============ Photo (PIL) ============
# 5 визуально разных изображений: цвет, форма, размер, формат.
# Намеренно разнообразные чтобы в склеенном Word'е было сразу видно где какое.

def _font(size):
    """Подбираем системный шрифт. Если ничего нет — fallback на default."""
    try:
        from PIL import ImageFont
        for name in ('arial.ttf', 'segoeui.ttf', 'DejaVuSans.ttf'):
            try:
                return ImageFont.truetype(name, size)
            except (OSError, IOError):
                pass
        return ImageFont.load_default()
    except Exception:
        return None


def make_photos():
    try:
        from PIL import Image, ImageDraw
    except ImportError:
        print('  (PIL не установлен — пропускаю фото; pip install Pillow)')
        return

    # 1. Логотип-плашка: тёмно-синий фон, белый круг с цифрой 1
    img = Image.new('RGB', (800, 600), (40, 60, 100))
    d = ImageDraw.Draw(img)
    d.ellipse([280, 180, 520, 420], fill=(240, 240, 240))
    f = _font(180)
    if f:
        d.text((360, 215), '1', fill=(40, 60, 100), font=f)
    f2 = _font(36)
    if f2:
        d.text((180, 470), 'photo-1-логотип', fill=(220, 220, 220), font=f2)
    img.save(TD / 'photo-1-логотип.png', 'PNG')
    print('  + photo-1-логотип.png')

    # 2. Bar chart: белый фон, цветные столбцы, подписи
    img = Image.new('RGB', (900, 500), (255, 255, 255))
    d = ImageDraw.Draw(img)
    bars = [(120, 320, 'Q1', (180, 60, 60)),
            (260, 200, 'Q2', (60, 180, 60)),
            (400, 150, 'Q3', (60, 100, 200)),
            (540, 240, 'Q4', (200, 160, 60))]
    base_y = 420
    for x, h, lbl, col in bars:
        d.rectangle([x, base_y - h, x + 100, base_y], fill=col)
        f3 = _font(28)
        if f3:
            d.text((x + 30, base_y + 10), lbl, fill=(40, 40, 40), font=f3)
    d.line([60, base_y, 700, base_y], fill=(80, 80, 80), width=2)
    f4 = _font(28)
    if f4:
        d.text((250, 30), 'Продажи 2024', fill=(20, 20, 20), font=f4)
    img.save(TD / 'photo-2-график.jpg', 'JPEG', quality=88)
    print('  + photo-2-график.jpg')

    # 3. Градиент: portrait, синий → оранжевый
    w, h = 600, 800
    img = Image.new('RGB', (w, h))
    px = img.load()
    for y in range(h):
        for x in range(w):
            t = (x + y) / (w + h)
            r = int(40 + (220 - 40) * t)
            g = int(80 + (140 - 80) * t)
            b = int(180 + (60 - 180) * t)
            px[x, y] = (r, g, b)
    d = ImageDraw.Draw(img)
    f5 = _font(54)
    if f5:
        d.text((140, 360), 'Градиент', fill=(255, 255, 255), font=f5)
    img.save(TD / 'photo-3-градиент.png', 'PNG')
    print('  + photo-3-градиент.png')

    # 4. Checkerboard pattern: квадратный, 5×5 клеток
    img = Image.new('RGB', (500, 500), (240, 240, 240))
    d = ImageDraw.Draw(img)
    cell = 100
    for r in range(5):
        for c in range(5):
            if (r + c) % 2 == 0:
                d.rectangle([c * cell, r * cell, (c + 1) * cell, (r + 1) * cell],
                            fill=(60, 120, 60))
    img.save(TD / 'photo-4-шашки.png', 'PNG')
    print('  + photo-4-шашки.png')

    # 5. Wide banner: тёмный градиент с большим текстом
    img = Image.new('RGB', (1600, 400), (20, 20, 30))
    d = ImageDraw.Draw(img)
    for x in range(1600):
        col = (20 + x // 40, 30 + x // 60, 60 + x // 30)
        d.line([x, 0, x, 400], fill=col)
    f6 = _font(120)
    if f6:
        d.text((250, 110), 'WIDE PHOTO', fill=(255, 240, 200), font=f6)
    img.save(TD / 'photo-5-баннер.jpg', 'JPEG', quality=90)
    print('  + photo-5-баннер.jpg')


# ============ CSV / TSV ============

def make_csv(filename, rows, delimiter=','):
    # Простая запись CSV/TSV. Кавычки используются если ячейка содержит
    # delimiter, кавычку или переводы строк.
    def esc_cell(v):
        s = '' if v is None else str(v)
        if delimiter in s or '"' in s or '\n' in s or '\r' in s:
            return '"' + s.replace('"', '""') + '"'
        return s
    lines = []
    for row in rows:
        lines.append(delimiter.join(esc_cell(c) for c in row))
    text = '\n'.join(lines) + '\n'
    (TD / filename).write_text(text, encoding='utf-8')
    print('  + ' + filename)


# ============ PDF (через reportlab) ============

def make_pdf_simple(filename, paragraphs):
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os as _os
    # Регистрируем системный шрифт с кириллицей (Arial из Windows)
    try:
        for font_path in ('C:/Windows/Fonts/arial.ttf', 'C:/Windows/Fonts/segoeui.ttf'):
            if _os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('Cyr', font_path))
                break
        else:
            pdfmetrics.registerFont(TTFont('Cyr', 'Helvetica'))
    except Exception:
        pass
    c = canvas.Canvas(str(TD / filename), pagesize=A4)
    c.setFont('Cyr', 12)
    w, h = A4
    margin = 50
    y = h - margin
    for para in paragraphs:
        # Перенос длинных строк ~80 символов на строку
        lines = []
        for ln in para.split('\n'):
            while len(ln) > 80:
                lines.append(ln[:80])
                ln = ln[80:]
            lines.append(ln)
        for ln in lines:
            if y < margin + 20:
                c.showPage()
                c.setFont('Cyr', 12)
                y = h - margin
            c.drawString(margin, y, ln)
            y -= 18
        y -= 10  # отступ между параграфами
    c.save()
    print('  + ' + filename)


def make_pdf_multipage(filename, pages):
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os as _os
    try:
        for font_path in ('C:/Windows/Fonts/arial.ttf', 'C:/Windows/Fonts/segoeui.ttf'):
            if _os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont('Cyr', font_path))
                break
    except Exception:
        pass
    c = canvas.Canvas(str(TD / filename), pagesize=A4)
    w, h = A4
    margin = 50
    for page_idx, paragraphs in enumerate(pages):
        c.setFont('Cyr', 14)
        c.drawString(margin, h - margin, f'Страница {page_idx + 1} из {len(pages)}')
        c.setFont('Cyr', 11)
        y = h - margin - 30
        for para in paragraphs:
            for ln in para.split('\n'):
                while len(ln) > 90:
                    if y < margin + 20: break
                    c.drawString(margin, y, ln[:90])
                    y -= 16
                    ln = ln[90:]
                if y < margin + 20: break
                c.drawString(margin, y, ln)
                y -= 16
            y -= 8
        if page_idx < len(pages) - 1:
            c.showPage()
    c.save()
    print('  + ' + filename)


# ============ Расширенные DOCX (через python-docx) ============

def make_docx_with_image(filename, image_path):
    from docx import Document
    from docx.shared import Cm
    doc = Document()
    doc.add_heading('Документ со встроенной картинкой', level=1)
    doc.add_paragraph('Этот Word содержит изображение прямо внутри документа. '
                      'Проверяет фикс B11 — file-merger должен скопировать его в выходной .docx.')
    doc.add_picture(str(image_path), width=Cm(10))
    doc.add_paragraph('Текст после картинки. Просто чтобы был контекст с двух сторон.')
    doc.save(str(TD / filename))
    print('  + ' + filename)


def make_docx_with_styles(filename):
    from docx import Document
    from docx.shared import Pt, RGBColor
    doc = Document()
    h1 = doc.add_heading('Раздел: Большой заголовок (H1)', level=1)
    doc.add_paragraph('Это обычный параграф под H1. Шрифт и интервалы должны сохраниться '
                      'в выходном Word\'е (тест фичи #5).')
    doc.add_heading('Подзаголовок (H2)', level=2)
    p = doc.add_paragraph()
    p.add_run('Это ').bold = False
    p.add_run('жирный текст').bold = True
    p.add_run(' посередине параграфа, и ')
    p.add_run('курсив').italic = True
    p.add_run('. Должны сохраниться в результате.')
    doc.add_heading('Список (для проверки numbering)', level=2)
    doc.add_paragraph('Первый пункт', style='List Number')
    doc.add_paragraph('Второй пункт', style='List Number')
    doc.add_paragraph('Третий пункт', style='List Number')
    doc.add_paragraph('Маркированный пункт', style='List Bullet')
    doc.add_paragraph('Ещё маркированный', style='List Bullet')
    # Цветной текст
    p = doc.add_paragraph()
    run = p.add_run('Цветной текст (красный)')
    run.font.color.rgb = RGBColor(0xC0, 0x40, 0x40)
    run.font.size = Pt(14)
    doc.save(str(TD / filename))
    print('  + ' + filename)


# ============ Расширенные XLSX (через openpyxl) ============

def make_xlsx_multisheet(filename):
    from openpyxl import Workbook
    wb = Workbook()
    # Лист 1: Сотрудники
    ws1 = wb.active
    ws1.title = 'Сотрудники'
    ws1.append(['Имя', 'Отдел', 'Зарплата'])
    ws1.append(['Иванов И.И.', 'IT', 80000])
    ws1.append(['Петров П.П.', 'Продажи', 65000])
    ws1.append(['Сидорова С.С.', 'Маркетинг', 70000])
    # Лист 2: Продажи
    ws2 = wb.create_sheet('Продажи')
    ws2.append(['Месяц', 'Сумма', 'Кол-во сделок'])
    ws2.append(['Январь', 1200000, 15])
    ws2.append(['Февраль', 1450000, 18])
    ws2.append(['Март', 1780000, 22])
    # Лист 3: Контакты
    ws3 = wb.create_sheet('Контакты клиентов')
    ws3.append(['Клиент', 'Телефон', 'Email'])
    ws3.append(['ООО Ромашка', '+7(495)111-22-33', 'romashka@example.ru'])
    ws3.append(['ИП Иванов', '+7(812)444-55-66', 'ivanov@example.ru'])
    wb.save(str(TD / filename))
    print('  + ' + filename)


def make_xlsx_merged(filename):
    from openpyxl import Workbook
    from openpyxl.styles import Alignment
    wb = Workbook()
    ws = wb.active
    ws.title = 'Отчёт'
    # Merged заголовок над двумя колонками
    ws['A1'] = 'Квартальный отчёт 2024'
    ws.merge_cells('A1:D1')
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    # Шапка таблицы
    ws.append([])  # empty row
    ws.append(['Месяц', 'Доход', 'Расход', 'Прибыль'])
    # Merged ячейка для Q1 (3 строки в первой колонке)
    ws.append(['Q1', 250000, 180000, 70000])
    ws.append([None, 280000, 195000, 85000])  # будет merged
    ws.append([None, 310000, 200000, 110000])  # будет merged
    ws.merge_cells('A4:A6')
    ws['A4'].alignment = Alignment(horizontal='center', vertical='center')
    # Аналогично Q2
    ws.append(['Q2', 295000, 210000, 85000])
    ws.append([None, 320000, 220000, 100000])
    ws.append([None, 340000, 230000, 110000])
    ws.merge_cells('A7:A9')
    ws['A7'].alignment = Alignment(horizontal='center', vertical='center')
    # Итого через merged
    ws.append(['Итого за полугодие', 1795000, 1235000, 560000])
    ws.merge_cells('A10:A10')  # тут не merged, но для теста
    wb.save(str(TD / filename))
    print('  + ' + filename)


# ============ Генерация ============

def main():
    print('== Word ==')
    make_docx('word-1-только-текст.docx',
        P('Это первый Word-документ с обычным текстом.') +
        P('Он содержит несколько параграфов для теста объединителя.') +
        P('Второй параграф — про погоду в Москве. Сейчас тут солнечно.') +
        P('И третий, чтобы было что объединять.')
    )
    make_docx('word-2-только-таблица.docx',
        TBL([
            ['Имя', 'Возраст', 'Город'],
            ['Иван', '30', 'Москва'],
            ['Мария', '25', 'Санкт-Петербург'],
            ['Алексей', '35', 'Казань'],
            ['Ольга', '28', 'Новосибирск'],
        ])
    )
    make_docx('word-3-текст-плюс-таблица.docx',
        P('Заголовок раздела «Параметры заказа»') +
        P('Это пример Word-документа со смешанным содержимым: сначала параграфы, потом таблица.') +
        TBL([
            ['Параметр', 'Значение'],
            ['Цена', '1500 ₽'],
            ['Количество', '5 шт.'],
            ['Скидка', '10%'],
            ['Итого', '6750 ₽'],
        ])
    )
    make_docx('word-4-несколько-таблиц.docx',
        P('Первая таблица — продажи по месяцам') +
        TBL([
            ['Месяц', 'Сумма'],
            ['Январь', '120 000 ₽'],
            ['Февраль', '145 000 ₽'],
            ['Март', '178 000 ₽'],
        ]) +
        P('Вторая таблица — расходы') +
        TBL([
            ['Категория', 'Бюджет'],
            ['Реклама', '30 000 ₽'],
            ['Зарплата', '80 000 ₽'],
            ['Аренда', '45 000 ₽'],
        ])
    )
    make_docx('word-5-длинный-текст.docx',
        P('Глава 1. Введение') +
        P('Здесь начинается описание проекта. Документ содержит три главы для демонстрации объединения длинных текстов.') +
        P('Каждая глава имеет несколько параграфов, чтобы при склейке было заметно, как тулзе удаётся сохранить структуру.') +
        P('Глава 2. Основная часть') +
        P('Описание основной задачи. Содержит несколько предложений с разными формулировками.') +
        P('Второй параграф главы 2 со списком: первое, второе, третье.') +
        P('Третий параграф главы 2 — продолжение мысли. Просто чтобы было больше контента.') +
        P('Глава 3. Заключение') +
        P('Краткие выводы по проекту. Все основные пункты раскрыты.') +
        P('Спасибо за внимание!')
    )

    print('== Excel ==')
    make_xlsx('excel-1-сотрудники.xlsx', 'Сотрудники', [
        ['Имя', 'Должность', 'Зарплата'],
        ['Иванов И.И.', 'Менеджер', 50000],
        ['Петров П.П.', 'Разработчик', 80000],
        ['Сидорова С.С.', 'Дизайнер', 65000],
        ['Смирнов А.А.', 'Аналитик', 75000],
    ])
    make_xlsx('excel-2-товары.xlsx', 'Прайс-лист', [
        ['№', 'Товар', 'Цена', 'Количество', 'Сумма'],
        [1, 'Хлеб белый', 50, 2, 100],
        [2, 'Молоко 3.2%', 80, 3, 240],
        [3, 'Сыр Российский', 350, 1, 350],
        [4, 'Яблоки', 120, 2, 240],
        [5, 'Огурцы', 90, 5, 450],
        [6, 'Помидоры', 150, 3, 450],
    ])
    make_xlsx('excel-3-финансы.xlsx', 'Бюджет', [
        ['Месяц', 'Доход', 'Расход', 'Прибыль'],
        ['Январь', 250000, 180000, 70000],
        ['Февраль', 280000, 195000, 85000],
        ['Март', 310000, 200000, 110000],
        ['Апрель', 295000, 210000, 85000],
    ])
    make_xlsx('excel-4-контакты.xlsx', 'Контакты', [
        ['ФИО', 'Телефон', 'Email', 'Город'],
        ['Иванов Иван Иванович', '+7(495)123-45-67', 'ivanov@example.ru', 'Москва'],
        ['Петрова Мария Алексеевна', '+7(812)987-65-43', 'petrova@example.ru', 'Санкт-Петербург'],
        ['Сидоров Алексей Константинович', '+7(383)555-12-34', 'sidorov@example.ru', 'Новосибирск'],
        ['Кузнецов Дмитрий Петрович', '+7(343)444-55-66', 'kuznetsov@example.ru', 'Екатеринбург'],
    ])
    make_xlsx('excel-5-расписание.xlsx', 'Расписание', [
        ['День', 'Время', 'Аудитория', 'Преподаватель', 'Дисциплина'],
        ['Понедельник', '09:00', '101', 'Иванов И.И.', 'Математика'],
        ['Понедельник', '10:30', '203', 'Петров П.П.', 'Физика'],
        ['Вторник', '09:00', '101', 'Иванов И.И.', 'Математика'],
        ['Вторник', '11:00', '305', 'Сидорова С.С.', 'История'],
        ['Среда', '09:00', '407', 'Кузнецов Д.П.', 'Информатика'],
    ])

    print('== Photos ==')
    make_photos()

    print('== CSV / TSV ==')
    make_csv('csv-1-простой.csv', [
        ['Имя', 'Возраст', 'Город'],
        ['Иван', 30, 'Москва'],
        ['Мария', 25, 'Санкт-Петербург'],
        ['Алексей', 35, 'Казань'],
    ])
    make_csv('csv-2-с-кавычками.csv', [
        ['Описание', 'Цена', 'Комментарий'],
        ['Книга "Война и мир"', 850, 'Классика, 4 тома'],
        ['Цитата: "Быть или не быть"', 0, 'Из Гамлета;\nШекспир'],
        ['Простой товар', 1200, 'Без специальных символов'],
    ])
    make_csv('tsv-1-табы.tsv', [
        ['Дата', 'Событие', 'Участников'],
        ['2024-01-15', 'Совещание отдела', 12],
        ['2024-02-01', 'Презентация продукта', 50],
        ['2024-03-10', 'Тренинг для новичков', 8],
    ], delimiter='\t')

    print('== PDF ==')
    make_pdf_simple('pdf-1-простой.pdf', [
        'Это простой PDF-документ с текстом на русском языке.',
        'Второй параграф для проверки извлечения текста через pdf.js.',
        'Третий параграф. Все три должны попасть в выходной Word отдельными абзацами.',
    ])
    make_pdf_multipage('pdf-2-многостраничный.pdf', [
        ['Содержимое первой страницы.', 'Несколько коротких параграфов для теста.'],
        ['Это вторая страница.', 'Между страницами PDF в выходном Word должен быть разрыв страницы.'],
        ['Третья и последняя страница.', 'Чтобы проверить что page-break срабатывает между всеми.'],
    ])
    make_pdf_simple('pdf-3-длинный.pdf', [
        'Глава 1. Введение в тему. ' * 5,
        'Раздел 1.1. Постановка задачи. ' * 4,
        'Раздел 1.2. Методология. ' * 4,
        'Глава 2. Основная часть. ' * 5,
        'Заключение. ' * 3,
    ])

    print('== Word (расширенные) ==')
    make_docx_with_image('word-6-с-картинкой.docx', TD / 'photo-1-логотип.png')
    make_docx_with_styles('word-7-со-стилями.docx')

    print('== Excel (расширенные) ==')
    make_xlsx_multisheet('excel-6-многолистовой.xlsx')
    make_xlsx_merged('excel-7-merged-cells.xlsx')

    print('---')
    print('Готово: файлы в ' + str(TD))


if __name__ == '__main__':
    main()
